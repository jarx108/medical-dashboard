"""
Generates the ControlManager Refactoring Review Word document.
Run: python generate_controlmanager_review.py
Output: ControlManager_Refactoring_Review.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x39, 0x64)
BLUE   = RGBColor(0x2E, 0x74, 0xB5)
TEAL   = RGBColor(0x00, 0x70, 0x70)
RED    = RGBColor(0xC0, 0x00, 0x00)
ORANGE = RGBColor(0xBF, 0x59, 0x00)
GREEN  = RGBColor(0x37, 0x86, 0x30)
GRAY   = RGBColor(0x59, 0x59, 0x59)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)
LIGHT_RED_HEX   = "FCEAEA"
LIGHT_GREEN_HEX = "EAF4EA"
LIGHT_BLUE_HEX  = "EBF3FB"

# ── Cell background helper ────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

# ── Style helpers ─────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = NAVY
    p.runs[0].font.size = Pt(18)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = BLUE
    p.runs[0].font.size = Pt(13)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = BLUE
    p.runs[0].font.size = Pt(11)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)

def body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(10.5)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def labelled(label, text, label_color=ORANGE):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}:  ")
    r1.bold = True
    r1.font.color.rgb = label_color
    r1.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)
    return p

def code_block(text, label=None):
    """Shaded monospace block, optional header label."""
    if label:
        lp = doc.add_paragraph()
        lp.paragraph_format.left_indent = Inches(0.3)
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(9)
        lr.font.color.rgb = GRAY
        lp.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name       = "Courier New"
    run.font.size       = Pt(8.5)
    run.font.color.rgb  = TEAL
    return p

def before_after(before_code, after_code):
    """Side-by-side 2-col table with Before / After code blocks."""
    t = doc.add_table(rows=2, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header cells
    headers = [("Before  ✗", "C00000"), ("After  ✓", "375E1E")]
    for ci, (label, color_hex) in enumerate(headers):
        cell = t.rows[0].cells[ci]
        set_cell_bg(cell, color_hex)
        run = cell.paragraphs[0].add_run(label)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9.5)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Code cells
    codes = [before_code, after_code]
    bg    = [LIGHT_RED_HEX, LIGHT_GREEN_HEX]
    for ci, (code, bg_hex) in enumerate(zip(codes, bg)):
        cell = t.rows[1].cells[ci]
        set_cell_bg(cell, bg_hex)
        p = cell.paragraphs[0]
        p.paragraph_format.left_indent = Inches(0.1)
        run = p.add_run(code)
        run.font.name      = "Courier New"
        run.font.size      = Pt(8)
        run.font.color.rgb = TEAL

    # Column widths
    for row in t.rows:
        row.cells[0].width = Inches(3.0)
        row.cells[1].width = Inches(3.0)

    doc.add_paragraph()

def make_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, "1F3964")
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        if ri % 2 == 1:
            for cell in tr.cells:
                set_cell_bg(cell, LIGHT_BLUE_HEX)
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def divider():
    p = doc.add_paragraph("─" * 90)
    p.runs[0].font.size = Pt(7)
    p.runs[0].font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
r = cover.add_run("ControlManager")
r.bold = True
r.font.size = Pt(28)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Refactoring Review & Recommendations")
r2.font.size = Pt(17)
r2.font.color.rgb = BLUE

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    "Project: Personal Medical Records Manager (prgPMR)",
    "File:    prgPMR/ControlManager.cs",
    "Review Date: February 2026",
]:
    r3 = meta.add_run(line + "\n")
    r3.font.size = Pt(11)
    r3.font.color.rgb = GRAY
    r3.font.name = "Courier New"

doc.add_paragraph()
summary_p = doc.add_paragraph()
summary_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = summary_p.add_run(
    "7 targeted changes  ·  eliminates all known OCP violations  ·  "
    "improves testability, safety, and clarity"
)
rs.italic = True
rs.font.size = Pt(10.5)
rs.font.color.rgb = GRAY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ORIGINAL CODE (annotated)
# ══════════════════════════════════════════════════════════════════════════════
h1("The Current ControlManager — Annotated")
body(
    "The listing below is the complete current source with problem markers. "
    "Each numbered annotation corresponds to a change described in the sections that follow."
)

code_block(
"""public class ControlManager
{
    public Button[] Buttons;                          // ❌ [2] public mutable field
    internal List<MedicalControl> MedicalControls { get; }
    private int activeControl = 0;
    private bool visible = false;

    public ControlManager(MainForm.MedicalControlType type, Button[] b)   // ❌ [1] depends on MainForm
    {
        MedicalControls = [];
        Buttons = b;
        switch (type)                                 // ❌ [1] knows every module type
        {
            case MainForm.MedicalControlType.Default:
                MedicalControls.Add(new DefaultControl(this));
                break;
            case MainForm.MedicalControlType.FamilyHistory:
                MedicalControls.Add(new FamilyHistoryControl(this));
                break;
            case MainForm.MedicalControlType.Medication:
                MedicalControls.Add(new MedicationsControl(this));
                break;
            case MainForm.MedicalControlType.Immunization:
                MedicalControls.AddRange(
                    new ImmunizationControl(this),
                    new ImmunizationDetailControl(this));
                break;
            case MainForm.MedicalControlType.DoctorVisit:
                MedicalControls.Add(new DoctorVisitsControl(this));
                break;
            // ... 4 more cases — every new module edits this file
            default:
                throw new ArgumentException("Invalid Medical Control type");
        }
    }

    public void ButtonClicked(int buttonIndex)        // ❌ [4] no bounds check, no null guard
    {
        MedicalControls[activeControl].ButtonActions[buttonIndex]();
    }

    public void RefreshVisibility()
    {
        SetVisible(visible);
    }

    public void SetVisible(bool isVisible)            // ❌ [6] no state reset on hide
    {
        visible = isVisible;
        foreach (MedicalControl m in MedicalControls)
        {
            if (isVisible && m == MedicalControls[activeControl])
            {
                m.Visible = true;
                for (int i = 0; i < Buttons.Length; i++)
                {
                    if (i >= m.ButtonsText.Length || m.ButtonsText[i] == null)
                        Buttons[i].Visible = false;
                    else
                    {
                        Buttons[i].Text = m.ButtonsText[i];
                        Buttons[i].Visible = true;
                    }
                }
            }
            else
            {
                m.Visible = false;
            }
        }
    }

    public bool GetVisible() { return visible; }      // ❌ [3] Java-style getter

    public void NextControl()
    {
        NextControl(null);                            // ❌ [5] null as magic reset signal
    }

    public void PreviousControl()
    {
        PreviousControl(null);                        // ❌ [5] null as magic reset signal
    }

    public void NextControl(DataInterface? data)
    {
        if (activeControl == MedicalControls.Count - 1)
            return;
        activeControl++;
        MedicalControls[activeControl].DataLoad(data);
        RefreshVisibility();
    }

    public void PreviousControl(DataInterface? data)  // ❌ [5] data on back-navigation is odd
    {
        if (activeControl == 0)
            return;
        activeControl--;
        MedicalControls[activeControl].DataLoad(data);
        RefreshVisibility();
    }
}"""
)

# ══════════════════════════════════════════════════════════════════════════════
# CHANGE 1 — SWITCH STATEMENT
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("Change 1 — Kill the Switch Statement")

labelled("Problem", "The constructor accepts MainForm.MedicalControlType and contains a switch "
         "over every module type. Every new module requires opening and editing ControlManager.cs. "
         "This is a direct violation of the Open/Closed Principle.", RED)

labelled("Root cause", "ControlManager is acting as both a factory (creating controls) and a "
         "coordinator (managing their lifecycle). These are two separate responsibilities.", ORANGE)

labelled("Fix", "Introduce IMedicalModule. Each module creates its own controls. "
         "ControlManager only receives the already-created list — it never needs to know "
         "what modules exist.", GREEN)

doc.add_paragraph()

before_after(
"""// Constructor coupled to MainForm enum
public ControlManager(
    MainForm.MedicalControlType type,
    Button[] b)
{
    MedicalControls = [];
    Buttons = b;
    switch (type)
    {
        case MedicalControlType.Immunization:
            MedicalControls.AddRange(
                new ImmunizationControl(this),
                new ImmunizationDetailControl(this));
            break;
        case MedicalControlType.DoctorVisit:
            MedicalControls.Add(
                new DoctorVisitsControl(this));
            break;
        // 7 more cases...
        default:
            throw new ArgumentException(
                "Invalid Medical Control type");
    }
}""",
"""// New IMedicalModule interface
public interface IMedicalModule
{
    string ModuleId    { get; }
    string DisplayName { get; }
    IReadOnlyList<MedicalControl>
        CreateControls(ControlManager m);
}

// Constructor — zero switch, zero module knowledge
public ControlManager(
    IMedicalModule module,
    Button[] buttons)
{
    _buttons = buttons;
    MedicalControls =
        [..module.CreateControls(this)];
}

// Each module wires its own controls:
public class ImmunizationModule : IMedicalModule
{
    public string ModuleId    => "immunization";
    public string DisplayName => "Immunizations";
    public IReadOnlyList<MedicalControl>
        CreateControls(ControlManager m) =>
        [ new ImmunizationControl(m),
          new ImmunizationDetailControl(m) ];
}"""
)

body("Adding a new module now requires zero changes to ControlManager.cs. "
     "One new class, registered in one line at startup.", italic=True)

# ══════════════════════════════════════════════════════════════════════════════
# CHANGE 2 — PUBLIC BUTTON FIELD
# ══════════════════════════════════════════════════════════════════════════════
divider()
h1("Change 2 — Make Buttons Private and Readonly")

labelled("Problem", "Buttons is a public mutable field. Any class anywhere in the project "
         "can reassign the array reference or mutate individual Button objects in it. "
         "The button bar is shared UI infrastructure — only ControlManager should drive it.", RED)

before_after(
"""public Button[] Buttons;   // anyone can write to this""",
"""private readonly Button[] _buttons;
// only SetVisible() ever touches _buttons"""
)

body("All internal usages already go through SetVisible(), so the only required "
     "follow-up is renaming Buttons → _buttons in that one method.", italic=True)

# ══════════════════════════════════════════════════════════════════════════════
# CHANGE 3 — GETVISIBLE PROPERTY
# ══════════════════════════════════════════════════════════════════════════════
divider()
h1("Change 3 — Replace GetVisible() with a Property")

labelled("Problem",
         "GetVisible() is a Java-style accessor method. C# convention is to use properties. "
         "It also creates inconsistency with how MedicalControl.Visible is already expressed "
         "as a property throughout the codebase.", ORANGE)

before_after(
"""public bool GetVisible()
{
    return visible;
}""",
"""public bool IsVisible => _visible;"""
)

body("Renamed to IsVisible to avoid any potential shadowing of "
     "the inherited Control.Visible property if the class hierarchy changes.", italic=True)

# ══════════════════════════════════════════════════════════════════════════════
# CHANGE 4 — BUTTONCLICKED BOUNDS CHECK
# ══════════════════════════════════════════════════════════════════════════════
divider()
h1("Change 4 — Add Bounds Checking to ButtonClicked")

labelled("Problem",
         "If buttonIndex is out of range for the active control's ButtonActions array, the "
         "call throws an IndexOutOfRangeException with no useful context. Additionally, "
         "ButtonActions slots are intentionally null for hidden buttons — calling a null "
         "action will throw a NullReferenceException.", RED)

labelled("Context",
         "The button bar always has 6 slots. A MedicalControl can expose fewer actions "
         "(e.g. [Add, null, null, Reset, null, null]). The null slots map to hidden buttons "
         "in SetVisible(), but ButtonClicked has no awareness of this.", ORANGE)

before_after(
"""public void ButtonClicked(int buttonIndex)
{
    // No bounds check — throws if index >= length
    // No null check — throws on null action slots
    MedicalControls[activeControl]
        .ButtonActions[buttonIndex]();
}""",
"""public void ButtonClicked(int buttonIndex)
{
    var actions =
        MedicalControls[_activeControl]
            .ButtonActions;

    if (buttonIndex < 0 ||
        buttonIndex >= actions.Length)
        return;

    // ?.Invoke() safely skips null slots
    // (null = button is hidden by design)
    actions[buttonIndex]?.Invoke();
}"""
)

# ══════════════════════════════════════════════════════════════════════════════
# CHANGE 5 — NULL MAGIC SIGNAL
# ══════════════════════════════════════════════════════════════════════════════
divider()
h1("Change 5 — Remove Null as a Magic Reset Signal")

labelled("Problem",
         "NextControl(null) and PreviousControl(null) pass null as a payload with the "
         "implicit meaning of 'reset/reload'. This forces every DataLoad override to "
         "check for null explicitly and interpret what null means in their context. "
         "It is an unclear contract that will silently break as the codebase grows.", RED)

labelled("Secondary issue",
         "PreviousControl(DataInterface? data) accepts a data payload for backwards "
         "navigation. Navigating back to a list view almost never needs to pass data — "
         "it just needs to reload. The overload design conflates two different intents.", ORANGE)

labelled("Fix",
         "Rename to NavigateForward/NavigateBack to make intent explicit, and give "
         "MedicalControl a dedicated Reload() method for the 'return to list' case.", GREEN)

before_after(
"""// null secretly means "please reset yourself"
public void NextControl()
{
    NextControl(null);
}
public void PreviousControl()
{
    PreviousControl(null);
}

public void NextControl(DataInterface? data)
{
    if (activeControl ==
        MedicalControls.Count - 1) return;
    activeControl++;
    // DataLoad must handle null itself
    MedicalControls[activeControl]
        .DataLoad(data);
    RefreshVisibility();
}

public void PreviousControl(DataInterface? data)
{
    if (activeControl == 0) return;
    activeControl--;
    MedicalControls[activeControl]
        .DataLoad(data);
    RefreshVisibility();
}""",
"""// Forward: always carries an explicit payload
public void NavigateForward(IControlPayload payload)
{
    if (_activeControl ==
        MedicalControls.Count - 1) return;
    _activeControl++;
    MedicalControls[_activeControl]
        .DataLoad(payload);
    RefreshVisibility();
}

// Back: no data needed — just reload the list
public void NavigateBack()
{
    if (_activeControl == 0) return;
    _activeControl--;
    MedicalControls[_activeControl]
        .Reload();    // explicit named intent
    RefreshVisibility();
}

// In MedicalControl (base class):
public virtual void DataLoad(
    IControlPayload payload) { }

public virtual void Reload() { }"""
)

body("IControlPayload replaces the empty DataInterface marker. "
     "See the Architecture Review document for the full IControlPayload definition.", italic=True)

# ══════════════════════════════════════════════════════════════════════════════
# CHANGE 6 — RESET ACTIVECONTROL ON HIDE
# ══════════════════════════════════════════════════════════════════════════════
divider()
h1("Change 6 — Reset activeControl When the Module is Hidden")

labelled("Problem",
         "When a user opens Immunization, navigates forward to the detail form, then clicks "
         "a different module in the nav panel without saving — SetVisible(false) is called "
         "but activeControl remains at 1 (the detail form). When the user returns to "
         "Immunization, they land directly on the detail form instead of the list view. "
         "This is unexpected and potentially shows stale unsaved data.", RED)

before_after(
"""public void SetVisible(bool isVisible)
{
    visible = isVisible;
    // activeControl is never reset —
    // user returns to wherever they left off,
    // even if that was a detail form mid-edit
    foreach (MedicalControl m in MedicalControls)
    {
        if (isVisible &&
            m == MedicalControls[activeControl])
        {
            m.Visible = true;
            // ... update buttons
        }
        else
        {
            m.Visible = false;
        }
    }
}""",
"""public void SetVisible(bool isVisible)
{
    _visible = isVisible;

    // Snap back to the list view (index 0)
    // whenever this module is hidden
    if (!isVisible && _activeControl != 0)
    {
        _activeControl = 0;
        MedicalControls[0].Reload();
    }

    foreach (MedicalControl m in MedicalControls)
    {
        if (isVisible &&
            m == MedicalControls[_activeControl])
        {
            m.Visible = true;
            for (int i = 0;
                 i < _buttons.Length; i++)
            {
                if (i >= m.ButtonsText.Length ||
                    m.ButtonsText[i] == null)
                    _buttons[i].Visible = false;
                else
                {
                    _buttons[i].Text =
                        m.ButtonsText[i];
                    _buttons[i].Visible = true;
                }
            }
        }
        else
        {
            m.Visible = false;
        }
    }
}"""
)

body("The Reload() call on index 0 also ensures the list grid is fresh "
     "when the user returns, picking up any records saved in other sessions.", italic=True)

# ══════════════════════════════════════════════════════════════════════════════
# CHANGE 7 — TESTABILITY
# ══════════════════════════════════════════════════════════════════════════════
divider()
h1("Change 7 — Expose ActiveControlIndex for Testability")

labelled("Problem",
         "The private _activeControl field is the primary state of ControlManager, but "
         "there is no way to assert on it in a unit test without using reflection. "
         "A read-only property costs nothing and makes the class fully testable.", ORANGE)

before_after(
"""private int activeControl = 0;
// No way to read this in a test
// without reflection""",
"""private int _activeControl = 0;

// Read-only — external code can observe
// but never mutate navigation state
public int ActiveControlIndex => _activeControl;"""
)

h3("Example Unit Tests Enabled by This Change")
code_block(
"""[Fact]
public void NavigateForward_AdvancesActiveIndex()
{
    var manager = new ControlManager(
        new ImmunizationModule(), fakeButtons);

    manager.NavigateForward(
        new ImmunizationPayload(ControlPayloadMode.Add));

    Assert.Equal(1, manager.ActiveControlIndex);
}

[Fact]
public void NavigateBack_DecrementsActiveIndex()
{
    var manager = BuildManagerAtIndex(1);
    manager.NavigateBack();
    Assert.Equal(0, manager.ActiveControlIndex);
}

[Fact]
public void NavigateForward_AtLastControl_DoesNotAdvance()
{
    var manager = BuildManagerWithOneControl();
    manager.NavigateForward(
        new ImmunizationPayload(ControlPayloadMode.Add));
    Assert.Equal(0, manager.ActiveControlIndex);
}

[Fact]
public void SetVisible_False_ResetsToIndex0()
{
    var manager = BuildManagerAtIndex(1);
    manager.SetVisible(false);
    Assert.Equal(0, manager.ActiveControlIndex);
}"""
)

# ══════════════════════════════════════════════════════════════════════════════
# FINAL — CLEANED UP CLASS
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("The Refactored ControlManager — Complete Listing")

body("All 7 changes applied. The class is now closed for modification when adding new modules, "
     "has no public mutable fields, uses C# conventions throughout, handles edge cases safely, "
     "and is fully unit-testable.")

code_block(
"""public class ControlManager
{
    // ── Fields ──────────────────────────────────────────────────────
    private readonly Button[] _buttons;
    private int  _activeControl = 0;
    private bool _visible       = false;

    // ── Properties ──────────────────────────────────────────────────
    internal IReadOnlyList<MedicalControl> MedicalControls { get; }
    public   bool IsVisible          => _visible;
    public   int  ActiveControlIndex => _activeControl;

    // ── Constructor ─────────────────────────────────────────────────
    public ControlManager(IMedicalModule module, Button[] buttons)
    {
        _buttons        = buttons;
        MedicalControls = [..module.CreateControls(this)];
    }

    // ── Button routing ──────────────────────────────────────────────
    public void ButtonClicked(int buttonIndex)
    {
        var actions = MedicalControls[_activeControl].ButtonActions;
        if (buttonIndex < 0 || buttonIndex >= actions.Length) return;
        actions[buttonIndex]?.Invoke();
    }

    // ── Navigation ──────────────────────────────────────────────────
    public void NavigateForward(IControlPayload payload)
    {
        if (_activeControl == MedicalControls.Count - 1) return;
        _activeControl++;
        MedicalControls[_activeControl].DataLoad(payload);
        RefreshVisibility();
    }

    public void NavigateBack()
    {
        if (_activeControl == 0) return;
        _activeControl--;
        MedicalControls[_activeControl].Reload();
        RefreshVisibility();
    }

    // ── Visibility ──────────────────────────────────────────────────
    public void RefreshVisibility() => SetVisible(_visible);

    public void SetVisible(bool isVisible)
    {
        _visible = isVisible;

        if (!isVisible && _activeControl != 0)
        {
            _activeControl = 0;
            MedicalControls[0].Reload();
        }

        foreach (MedicalControl m in MedicalControls)
        {
            if (isVisible && m == MedicalControls[_activeControl])
            {
                m.Visible = true;
                for (int i = 0; i < _buttons.Length; i++)
                {
                    if (i >= m.ButtonsText.Length || m.ButtonsText[i] == null)
                        _buttons[i].Visible = false;
                    else
                    {
                        _buttons[i].Text    = m.ButtonsText[i];
                        _buttons[i].Visible = true;
                    }
                }
            }
            else
            {
                m.Visible = false;
            }
        }
    }
}"""
)

# ══════════════════════════════════════════════════════════════════════════════
# SUMMARY TABLE
# ══════════════════════════════════════════════════════════════════════════════
h1("Summary of All Changes")

make_table(
    ["#", "Change", "Category", "Impact"],
    [
        ["1", "Replace switch + enum with IMedicalModule",           "OCP / SRP",       "New modules never touch this file"],
        ["2", "Make Buttons a private readonly field",               "Encapsulation",   "Prevents external mutation of shared UI"],
        ["3", "Replace GetVisible() with IsVisible property",        "Convention",      "Consistent C# idiom, cleaner call sites"],
        ["4", "Add bounds check + null-safe invoke to ButtonClicked","Safety",          "No more IndexOutOfRange or NullReference"],
        ["5", "Replace null magic with NavigateForward/NavigateBack","Clarity / API",   "Explicit intent, removes ambiguous null contract"],
        ["6", "Reset activeControl in SetVisible(false)",            "UX correctness",  "User always returns to list view, not mid-edit detail"],
        ["7", "Expose ActiveControlIndex read-only property",        "Testability",     "Full unit test coverage without reflection"],
    ],
    col_widths=[0.3, 2.4, 1.3, 2.4]
)

body(
    "These changes leave the core navigation logic (ControlManager coordinating "
    "MedicalControl sub-panels via DataLoad/Reload) entirely intact. They address "
    "only what is wrong, without redesigning what is already right.",
    italic=True
)

# ── Save ─────────────────────────────────────────────────────────────────────
output_path = "E:/Dev/GitHub/medical-dashboard/ControlManager_Refactoring_Review.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
