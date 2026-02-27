"""
Generates the Senior Architect Review Word document for prgPMR.
Run: python generate_review.py
Output: PMR_Architecture_Review.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.15)
    section.right_margin  = Inches(1.15)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x39, 0x64)   # headings
BLUE   = RGBColor(0x2E, 0x74, 0xB5)   # sub-headings
TEAL   = RGBColor(0x00, 0x70, 0x70)   # code / mono
RED    = RGBColor(0xC0, 0x00, 0x00)   # critical
ORANGE = RGBColor(0xBF, 0x59, 0x00)   # high severity
GRAY   = RGBColor(0x59, 0x59, 0x59)   # table text
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)

# ── Table header background helper ────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_bg_color(cell, r, g, b):
    hex_color = f"{r:02X}{g:02X}{b:02X}"
    set_cell_bg(cell, hex_color)

# ── Style helpers ─────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = NAVY
    p.runs[0].font.size = Pt(18)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = BLUE
    p.runs[0].font.size = Pt(14)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = BLUE
    p.runs[0].font.size = Pt(12)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)

def body(text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.bold   = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def numbered(text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def code_block(text):
    """Monospace shaded paragraph for code / diagrams."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.3)
    p.paragraph_format.right_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    # light-gray background
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = TEAL
    return p

def info_box(label, text, color=ORANGE):
    """Bold coloured label + text on same paragraph."""
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.color.rgb = color
    r1.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(3)
    return p

def severity_bullet(sev: str, text: str):
    colors = {"Critical": RED, "High": ORANGE, "Medium": RGBColor(0xBF,0x8F,0x00), "Low": GRAY}
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(f"[{sev}] ")
    r1.bold = True
    r1.font.color.rgb = colors.get(sev, GRAY)
    r1.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    return p

def make_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg_color(cell, 0x1F, 0x39, 0x64)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        if ri % 2 == 1:
            for cell in tr.cells:
                set_cell_bg(cell, "EBF3FB")
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
r = cover.add_run("Personal Medical Records Manager")
r.bold = True
r.font.size = Pt(26)
r.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("Senior Architect Review")
r2.font.size = Pt(18)
r2.font.color.rgb = BLUE

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    "Platform: Windows Forms / .NET 9.0-windows7.0",
    "Database: SQL Server (Microsoft.Data.SqlClient)",
    "Project State: Early Development — Immunization module partially implemented",
    "Review Date: February 2026",
]:
    r3 = meta.add_run(line + "\n")
    r3.font.size = Pt(11)
    r3.font.color.rgb = GRAY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. HIGH-LEVEL ARCHITECTURE OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1("1. High-Level Architecture Overview")

h2("System Purpose")
body(
    "An offline, single-user Windows desktop application for managing personal medical records. "
    "The user authenticates locally, then navigates between medical modules (Immunizations, Doctor "
    "Visits, Medications, etc.) through a shell form with a side navigation panel and a shared "
    "dynamic button bar at the bottom."
)

h2("Current Layer Map")
code_block(
"""┌─────────────────────────────────────────────────────────────┐
│                    PRESENTATION LAYER                        │
│  MainForm (shell/nav host)  │  LoginForm (auth)             │
│  MedicalControl subclasses (UserControls embedded in shell) │
├─────────────────────────────────────────────────────────────┤
│                   COORDINATION LAYER                         │
│  ControlManager (lifecycle, visibility, navigation)          │
│  MedicalControl (abstract base – buttons + DataLoad)        │
├─────────────────────────────────────────────────────────────┤
│                     DATA TRANSFER                            │
│  DataInterface (marker)  ImmunizationAddData  EditData      │
├─────────────────────────────────────────────────────────────┤
│                     DATA LAYER (missing)                     │
│  Planned SQL Server access — currently inline / stubbed      │
└─────────────────────────────────────────────────────────────┘"""
)

h2("Data Flow (As-Built)")
code_block(
"""User clicks nav button
  → MainForm.DisplayControl_Click
    → MainForm.DisplayMedicalControl(type)
      → ControlManager.SetVisible(true/false) for all managers
        → MedicalControl.Visible = true/false
          → Button bar text/visibility updated

User clicks action button (Add/Edit/Delete/Reset/Save/Cancel)
  → MainForm.ButtonBar_Click
    → ControlManagerDict[ActiveControl].ButtonClicked(index)
      → MedicalControls[activeControl].ButtonActions[index]()
        → e.g. ImmunizationControl.Add()
          → Manager.NextControl(new ImmunizationAddData(...))
            → activeControl++
            → ImmunizationDetailControl.DataLoad(data)
            → ControlManager.RefreshVisibility()"""
)

# ══════════════════════════════════════════════════════════════════════════════
# 2. MODULE-BY-MODULE ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Module-by-Module Analysis")

# MainForm
h2("MainForm — Navigation Shell")
make_table(
    ["Aspect", "Assessment"],
    [
        ["Responsibility", "Hosts panels, routes nav clicks, owns shared button bar"],
        ["Boundary Violation", "Owns MedicalControlType enum — module identity belongs to the module system, not the shell"],
        ["Dependencies", "Directly references every ControlManager and every MedicalControlType"],
        ["SOLID", "Violates OCP: adding a module requires editing MainForm to add a button and an enum value"],
        ["Issue", "Hardcoded user data (txtLastName.Text = \"Jangaon\") in constructor"],
    ],
    col_widths=[1.6, 4.8]
)

# ControlManager
h2("ControlManager — Module Lifecycle Coordinator")
make_table(
    ["Aspect", "Assessment"],
    [
        ["Responsibility", "Creates controls, tracks active sub-panel, drives visibility + button bar"],
        ["Boundary Violation", "Owns a switch statement over all module types — the worst OCP violation in the project"],
        ["Dependencies", "Knows about every concrete control class"],
        ["SOLID", "Violates OCP, partially violates SRP (creates + manages + routes)"],
        ["Good Pattern", "NextControl(DataInterface?) / PreviousControl() navigation is well-conceived"],
        ["Issue", "Buttons field is public — exposes shared UI infrastructure directly"],
    ],
    col_widths=[1.6, 4.8]
)

# MedicalControl
h2("MedicalControl — Abstract Module Base")
make_table(
    ["Aspect", "Assessment"],
    [
        ["Responsibility", "Declares button contract, passes ControlManager reference down"],
        ["Design", "The button preset dictionary (LowerbuttonBarPresetGrouping) is a smart reuse pattern"],
        ["Issue", "Each concrete control duplicates its own lowerbuttonBarPresetActionDict — should live in base"],
        ["Issue", "DataLoad(DataInterface?) is virtual with an empty default — subclasses can silently ignore it"],
        ["SOLID", "Reasonably good SRP; LSP would break if subclasses silently no-op DataLoad"],
    ],
    col_widths=[1.6, 4.8]
)

# ImmunizationControl
h2("ImmunizationControl — Most Complete Module")
make_table(
    ["Aspect", "Assessment"],
    [
        ["Good", "Single/double-click disambiguation via Timer is the correct WinForms technique"],
        ["Good", "Uses DataInterface typed dispatch in DataLoad"],
        ["Issue", "Hardcoded DataTable rows — no data layer exists yet"],
        ["Issue", "Delete() resets to GridInitial BEFORE the confirmation dialog — buttons flash even on 'No'"],
        ["Issue", "DataLoad(null) reinitialises the grid — fragile null-semantics"],
    ],
    col_widths=[1.6, 4.8]
)

# ImmunizationDetailControl
h2("ImmunizationDetailControl — Form Detail View")
make_table(
    ["Aspect", "Assessment"],
    [
        ["Good", "Dirty-state tracking (isUserDataControlsModified) is correct UX"],
        ["Good", "AttachEventHandlerstoUserDataControls is a structured approach"],
        ["Bug", "Cancel() does nothing when the form is clean — user expects navigation back regardless"],
        ["Issue", "FillAllUserControls() fills a single hardcoded label — not wired to real data"],
        ["Issue", "Save() shows a placeholder message box — no persistence implemented"],
    ],
    col_widths=[1.6, 4.8]
)

# LoginForm
h2("LoginForm — Authentication")
make_table(
    ["Aspect", "Assessment"],
    [
        ["Critical Bug", "pgmPMRMain.cs launches new MainForm() directly — login is completely bypassed"],
        ["Security", "Connection string hardcoded with machine-specific server name (EJVIDSYS\\SQLEXPRESS)"],
        ["Security", "Malformed keyword 'GridInitial Catalog' in connection string (extra word 'Grid')"],
        ["Security", "Password comparison with SELECT * implies plaintext password storage in DB"],
        ["Design", "LoginForm directly instantiates MainForm — tight coupling between auth and shell"],
        ["Pattern", "Should use ExecuteScalar, not SqlDataAdapter.Fill, for authentication queries"],
    ],
    col_widths=[1.6, 4.8]
)

# DataObjects
h2("DataObjects/ — Data Transfer Layer")
make_table(
    ["Aspect", "Assessment"],
    [
        ["DataInterface", "Empty marker interface — no contract enforced"],
        ["ImmunizationAddData", "Has SampleText = text — placeholder, not domain-representative"],
        ["ImmunizationEditData", "Completely empty — provides no data to the edit form"],
        ["Coverage", "Only Immunization has transfer objects — all other modules have none"],
    ],
    col_widths=[1.6, 4.8]
)

# Stubs
h2("Stub Modules (FamilyHistory, Medications, DoctorVisits, etc.)")
body(
    "All stub modules call SetButtons([], []) and contain no logic. "
    "DoctorVisitsDetail has an inconsistent name — it should be DoctorVisitsDetailControl "
    "to match the naming convention used by every other module pair."
)

# ══════════════════════════════════════════════════════════════════════════════
# 3. EXTENSIBILITY EVALUATION
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Extensibility Evaluation")

h2("Current Friction — Adding a New Module Today Requires:")
for step in [
    "Add a value to the MainForm.MedicalControlType enum",
    "Add a navigation button + event handler in MainForm.Designer.cs",
    "Add a case in the ControlManager switch statement",
    "Create concrete MedicalControl subclass(es)",
    "Add DataInterface implementations in DataObjects/",
]:
    numbered(step)

body(
    "This is 5 separate changes across 3 unrelated files for every new module — "
    "a textbook Open/Closed Principle failure.",
    bold=True
)

h2("What Works Well")
bullet("The MedicalControl abstraction is the right seam — it isolates module rendering from the shell")
bullet("DataInterface + typed dispatch pattern in DataLoad is clean and extensible as-is")
bullet("LowerbuttonBarPresetGrouping enum allows reuse of button bar semantics across all modules")

# ══════════════════════════════════════════════════════════════════════════════
# 4. DESIGN ISSUES INVENTORY
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Design Issues Inventory")

make_table(
    ["#", "Issue", "Severity", "Location"],
    [
        ["1",  "Login completely bypassed in Main()",                            "Critical", "pgmPMRMain.cs:30"],
        ["2",  "Malformed connection string keyword 'GridInitial Catalog'",       "Critical", "LoginForm.cs:28"],
        ["3",  "Plaintext password comparison in SQL",                           "Critical", "LoginForm.cs:80"],
        ["4",  "MedicalControlType enum defined inside MainForm",                "High",     "MainForm.cs:6"],
        ["5",  "ControlManager switch statement over all module types",          "High",     "ControlManager.cs:21"],
        ["6",  "All module types hardcoded — zero plugin capability",            "High",     "System-wide"],
        ["7",  "No data access layer / repository pattern",                      "High",     "System-wide"],
        ["8",  "Hardcoded user info in MainForm constructor",                    "Medium",   "MainForm.cs:95"],
        ["9",  "Cancel() does nothing on clean (unmodified) form",               "Medium",   "ImmunizationDetailControl.cs:129"],
        ["10", "Delete() resets buttons before user confirms",                   "Medium",   "ImmunizationControl.cs:70"],
        ["11", "DataLoad(null) used as magic reset signal",                      "Medium",   "ImmunizationControl.cs:48"],
        ["12", "DataInterface is an empty marker with no enforced contract",     "Medium",   "DataInterface.cs"],
        ["13", "ControlManager.Buttons is public",                               "Low",      "ControlManager.cs:12"],
        ["14", "DoctorVisitsDetail vs *DetailControl naming inconsistency",      "Low",      "DoctorVisitsDetail.cs"],
        ["15", "TestsControl uses label text to confirm actions (debug code)",   "Low",      "TestsControl.cs:25"],
        ["16", "Empty auto-generated event handlers committed",                  "Low",      "SurgeriesDetailControl.cs:20"],
        ["17", "UpdateUserInfoControl calls no SetButtons",                      "Low",      "UpdateUserInfoControl.cs:17"],
        ["18", "Each detail control duplicates lowerbuttonBarPresetActionDict",  "Low",      "All detail controls"],
    ],
    col_widths=[0.3, 3.2, 0.9, 2.0]
)

# ══════════════════════════════════════════════════════════════════════════════
# 5. PROPOSED DOMAIN MODEL
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Proposed Domain Model")

body(
    "The following model provides strongly-typed, immutable domain objects that scale across all "
    "current and future medical modules. Every record inherits from MedicalRecord, which carries "
    "identity, patient linkage, audit timestamps, and attachment support."
)

code_block(
"""// ── Core patient identity ─────────────────────────────────────
public record PatientProfile(
    Guid Id, string FirstName, string LastName,
    DateOnly DateOfBirth, string? BloodType);

// ── Shared value objects ──────────────────────────────────────
public record ProviderRef(string Name, string? Facility, string? Phone);
public record DoseRecord(DateOnly Date, string? LotNumber, string? Site);
public record Attachment(Guid Id, string FileName, AttachmentType Type, byte[] Data);
public enum AttachmentType { Pdf, Image }

// ── Shared base ───────────────────────────────────────────────
public abstract class MedicalRecord
{
    public Guid Id           { get; init; } = Guid.NewGuid();
    public Guid PatientId    { get; init; }
    public DateTimeOffset CreatedAt  { get; init; } = DateTimeOffset.UtcNow;
    public DateTimeOffset? ModifiedAt { get; set; }
    public IReadOnlyList<Attachment> Attachments { get; init; } = [];
}

// ── Module: Immunizations ─────────────────────────────────────
public class Immunization : MedicalRecord
{
    public VaccineGroup Group       { get; init; }
    public string VaccineName       { get; init; }
    public string? Manufacturer     { get; init; }
    public IReadOnlyList<DoseRecord> Doses { get; init; }
    public string? Notes            { get; init; }
}

// ── Module: Doctor Visits ─────────────────────────────────────
public class DoctorVisit : MedicalRecord
{
    public DateOnly VisitDate        { get; init; }
    public ProviderRef Provider      { get; init; }
    public string Reason             { get; init; }
    public string? Diagnosis         { get; init; }
    public string? Notes             { get; init; }
    public IReadOnlyList<Prescription> Prescriptions { get; init; }
}

// ── Module: Medications ───────────────────────────────────────
public class Medication : MedicalRecord
{
    public string DrugName           { get; init; }
    public string Dosage             { get; init; }
    public string Frequency          { get; init; }
    public DateOnly StartDate        { get; init; }
    public DateOnly? EndDate         { get; init; }
    public ProviderRef PrescribedBy  { get; init; }
}

// ── Module: Surgeries ─────────────────────────────────────────
public class Surgery : MedicalRecord
{
    public DateOnly SurgeryDate      { get; init; }
    public string Procedure          { get; init; }
    public ProviderRef Surgeon       { get; init; }
    public string? Hospital          { get; init; }
    public string? Notes             { get; init; }
}"""
)

# ══════════════════════════════════════════════════════════════════════════════
# 6. REFACTORING RECOMMENDATIONS
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Refactoring Recommendations")

h2("6a. Introduce a Module Registry — Eliminates the Switch Statement")
body(
    "Replace the MainForm.MedicalControlType enum and ControlManager switch with a "
    "self-registering IMedicalModule interface. Adding a new module then requires "
    "zero changes to existing files."
)
code_block(
"""public interface IMedicalModule
{
    string ModuleId     { get; }   // "immunization", "doctor-visits"
    string DisplayName  { get; }   // "Immunizations"
    IReadOnlyList<MedicalControl> CreateControls(ControlManager manager);
}

public class ImmunizationModule : IMedicalModule
{
    public string ModuleId    => "immunization";
    public string DisplayName => "Immunizations";
    public IReadOnlyList<MedicalControl> CreateControls(ControlManager m) =>
        [new ImmunizationControl(m), new ImmunizationDetailControl(m)];
}

public class ModuleRegistry
{
    private readonly List<IMedicalModule> _modules = [];
    public void Register(IMedicalModule module) => _modules.Add(module);
    public IEnumerable<IMedicalModule> All => _modules;
}

// ControlManager becomes:
public ControlManager(IMedicalModule module, Button[] buttons)
{
    MedicalControls = [..module.CreateControls(this)];
    Buttons = buttons;
}"""
)

h2("6b. Repository Pattern for Data Access")
code_block(
"""public interface IRepository<T> where T : MedicalRecord
{
    Task<IReadOnlyList<T>> GetAllAsync(Guid patientId);
    Task<T?> GetByIdAsync(Guid id);
    Task AddAsync(T record);
    Task UpdateAsync(T record);
    Task DeleteAsync(Guid id);
}

// Each MedicalControl receives its repository via constructor injection:
public class ImmunizationControl : MedicalControl
{
    private readonly IRepository<Immunization> _repo;
    public ImmunizationControl(ControlManager m, IRepository<Immunization> repo)
        : base(m) => _repo = repo;
}"""
)

h2("6c. Give DataInterface a Real Contract")
code_block(
"""public interface IControlPayload
{
    ControlPayloadMode Mode { get; }
}

public enum ControlPayloadMode { Add, Edit, View }

public record ImmunizationPayload(ControlPayloadMode Mode, Immunization? Record = null)
    : IControlPayload;"""
)

h2("6d. Move Preset Action Dictionary to MedicalControl Base")
code_block(
"""// In MedicalControl — each subclass only provides its actions array
protected void SetButtonsFromPreset(
    LowerbuttonBarPresetGrouping grouping,
    Action[] actions)
{
    SetButtons(lowerbuttonBarPresetTextsDict[grouping], actions);
}"""
)

h2("6e. Fix the Cancel Bug")
code_block(
"""public void Cancel()
{
    if (isUserDataControlsModified)
    {
        var result = MessageBox.Show(
            "Data has been modified. Cancel without saving?",
            "Confirmation", MessageBoxButtons.YesNo, MessageBoxIcon.Question);
        if (result != DialogResult.Yes) return;
    }
    Manager.PreviousControl(); // always navigate back
}"""
)

h2("6f. Fix Login Entry Point")
code_block(
"""// pgmPMRMain.cs — launch LoginForm, NOT MainForm
static void Main()
{
    Application.EnableVisualStyles();
    Application.SetCompatibleTextRenderingDefault(false);
    Application.Run(new LoginForm()); // was: new MainForm()
}"""
)

h2("6g. Recommended Folder Structure")
code_block(
"""prgPMR/
├── Core/
│   ├── MedicalControl.cs
│   ├── ControlManager.cs
│   ├── ModuleRegistry.cs
│   └── Abstractions/
│       ├── IMedicalModule.cs
│       ├── IControlPayload.cs
│       └── AbstractControlDescriptionProvider.cs
├── Domain/
│   ├── PatientProfile.cs
│   ├── MedicalRecord.cs
│   ├── Immunization.cs
│   ├── DoctorVisit.cs
│   ├── Medication.cs
│   ├── Surgery.cs
│   ├── Hospitalization.cs
│   └── Shared/
│       ├── ProviderRef.cs
│       ├── DoseRecord.cs
│       └── Attachment.cs
├── Data/
│   ├── IRepository.cs
│   ├── SqlConnectionFactory.cs
│   ├── Sql/
│   │   ├── SqlImmunizationRepository.cs
│   │   └── SqlDoctorVisitRepository.cs
│   └── Stub/
│       └── InMemoryImmunizationRepository.cs
├── Modules/
│   ├── Immunization/
│   │   ├── ImmunizationModule.cs
│   │   ├── ImmunizationControl.cs
│   │   ├── ImmunizationDetailControl.cs
│   │   └── ImmunizationPayload.cs
│   ├── DoctorVisits/
│   │   └── ...
│   └── Medications/
│       └── ...
├── Shell/
│   ├── MainForm.cs
│   ├── LoginForm.cs
│   └── DefaultControl.cs
└── pgmPMRMain.cs"""
)

# ══════════════════════════════════════════════════════════════════════════════
# 7. DATA FLOW FOR KEY OPERATIONS
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Data Flow for Key Operations")

h2("7a. Adding a Doctor Visit (Target State)")
code_block(
"""1. User clicks "Doctor Visits" in nav panel
   → MainForm.DisplayControl_Click
   → ControlManager["doctor-visits"].SetVisible(true)
   → DoctorVisitsControl shown, buttons = [Add, null, null, Reset]

2. User clicks "Add"
   → MainForm.ButtonBar_Click(index=0)
   → ControlManager.ButtonClicked(0)
   → DoctorVisitsControl.Add()
   → Manager.NextControl(new DoctorVisitPayload(Mode: Add))

3. ControlManager.NextControl(payload)
   → activeControl = 1  (DoctorVisitsDetailControl)
   → DoctorVisitsDetailControl.DataLoad(payload)
     → SetButtonsFromPreset(DetailAdd, [...])
     → ClearAllFields()

4. User fills form, clicks "Save"
   → MainForm.ButtonBar_Click(index=4)
   → DoctorVisitsDetailControl.Save()
   → _repo.AddAsync(BuildDomainObject())     ← async call
   → Manager.PreviousControl()
   → DoctorVisitsControl.DataLoad(null)      ← refreshes grid

5. Grid repopulates from repository"""
)

h2("7b. Save / Load Records")
code_block(
"""Application startup:
  pgmPMRMain → LoginForm → validates credentials via SQL
  → Creates SqlConnectionFactory with validated connection string
  → Passes factory to each module's repository
  → MainForm initialised with ModuleRegistry

Module open (first time):
  ControlManager.SetVisible(true)
  → MedicalControl.DataLoad(null)          ← null = "reload from source"
  → await _repo.GetAllAsync(patientId)
  → Bind result to DataGridView

Record save:
  DetailControl.Save()
  → Validate fields (local)
  → Build domain object
  → await _repo.AddAsync(record)           ← throws on DB error
  → Show success / navigate back"""
)

h2("7c. Module Initialisation (Target State)")
code_block(
"""pgmPMRMain.Main()
  → ModuleRegistry registry = new()
  → registry.Register(new ImmunizationModule())
  → registry.Register(new DoctorVisitModule())
  → registry.Register(new AllergiesModule())   // just add this line
  → var factory = new SqlConnectionFactory(configuredConnString)
  → Application.Run(new LoginForm(registry, factory))

LoginForm.OnLoginSuccess()
  → MainForm mainForm = new(registry, factory, authenticatedUser)
  → mainForm.Show(); this.Hide()

MainForm constructor:
  → foreach module in registry.All:
      ControlManager m = new(module, lowerButtonBar)
      ControlManagerDict[module.ModuleId] = m
      foreach control in m.MedicalControls:
          pnlMain.Controls.Add(control)
      Build nav button dynamically from module.DisplayName"""
)

# ══════════════════════════════════════════════════════════════════════════════
# 8. DIAGRAMS
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Architecture Diagrams")

h2("8a. Overall Architecture")
code_block(
"""┌──────────────────────────────────────────────────────────────────────┐
│                          Shell Layer                                  │
│   ┌──────────────┐       ┌──────────────────────────────────────┐   │
│   │  LoginForm   │──────▶│              MainForm                │   │
│   │  (auth)      │       │  nav panel │ pnlMain │ button bar    │   │
│   └──────────────┘       └──────────────────────────────────────┘   │
└──────────────────────────────────────────────────────────────────────┘
                                     │ hosts N ControlManagers
                                     ▼
┌──────────────────────────────────────────────────────────────────────┐
│                        Coordination Layer                             │
│   ┌──────────────────────────────────────────────────────────────┐  │
│   │                    ControlManager [1..N]                      │  │
│   │  activeControl idx │ SetVisible │ NextControl │ PrevControl   │  │
│   └──────────────────────────────────────────────────────────────┘  │
│                   │ owns List<MedicalControl>                        │
└──────────────────────────────────────────────────────────────────────┘
                                     │
                                     ▼
┌──────────────────────────────────────────────────────────────────────┐
│                         Module Layer                                  │
│   ┌─────────────────────┐   ┌──────────────────────┐               │
│   │  MedicalControl     │──▶│  MedicalControl       │  (abstract)  │
│   │  (list/grid view)   │   │  (detail/edit view)   │              │
│   └─────────────────────┘   └──────────────────────┘               │
│      ImmunizationControl        ImmunizationDetailControl           │
│      DoctorVisitsControl        DoctorVisitsDetailControl           │
└──────────────────────────────────────────────────────────────────────┘
                                     │ reads / writes
                                     ▼
┌──────────────────────────────────────────────────────────────────────┐
│                          Data Layer                                   │
│   IRepository<T>                                                     │
│   SqlImmunizationRepository    InMemoryImmunizationRepository        │
│   SqlDoctorVisitRepository     ...                                   │
│                              │                                       │
│                         SQL Server (local)                           │
└──────────────────────────────────────────────────────────────────────┘"""
)

h2("8b. Class Diagram (Mermaid)")
code_block(
"""classDiagram
    class MainForm {
        -Dict~string,ControlManager~ ControlManagerDict
        -string ActiveModuleId
        +DisplayMedicalControl(moduleId)
        +ButtonBar_Click(sender, e)
        +DisplayControl_Click(sender, e)
    }

    class ControlManager {
        -int activeControl
        -bool visible
        +List~MedicalControl~ MedicalControls
        +ButtonClicked(int index)
        +SetVisible(bool)
        +NextControl(IControlPayload?)
        +PreviousControl(IControlPayload?)
    }

    class MedicalControl {
        <<abstract>>
        #ControlManager Manager
        +string[] ButtonsText
        +Action[] ButtonActions
        +DataLoad(IControlPayload?)
        #SetButtons(string[], Action[])
    }

    class IMedicalModule {
        <<interface>>
        +string ModuleId
        +string DisplayName
        +CreateControls(ControlManager) List~MedicalControl~
    }

    class IRepository~T~ {
        <<interface>>
        +GetAllAsync(Guid) Task~List~T~~
        +AddAsync(T)       Task
        +UpdateAsync(T)    Task
        +DeleteAsync(Guid) Task
    }

    MainForm "1" --> "N" ControlManager
    ControlManager "1" --> "1..N" MedicalControl
    MedicalControl <|-- ImmunizationControl
    MedicalControl <|-- ImmunizationDetailControl
    IMedicalModule <|.. ImmunizationModule
    ImmunizationControl --> IRepository"""
)

h2("8c. Sequence — Adding an Immunization Record")
code_block(
"""User       MainForm    ControlManager   ImmunControl   ImmunDetailControl   Repository
 │               │               │               │               │               │
 │ click [Add]   │               │               │               │               │
 │──────────────▶│               │               │               │               │
 │               │ ButtonBar_Click               │               │               │
 │               │──────────────▶│               │               │               │
 │               │               │ ButtonClicked(0)              │               │
 │               │               │──────────────▶│               │               │
 │               │               │               │ Add()         │               │
 │               │               │ NextControl(Payload{Add})     │               │
 │               │               │◀──────────────│               │               │
 │               │               │ activeControl++               │               │
 │               │               │ DataLoad(payload)────────────▶│               │
 │               │               │               │               │SetButtons()   │
 │               │               │               │               │ClearFields()  │
 │               │               │ RefreshVisibility             │               │
 │ [form shown]  │ button bar updates             │               │               │
 │               │               │               │               │               │
 │ fill & click [Save]           │               │               │               │
 │──────────────▶│               │               │               │               │
 │               │ ButtonBar_Click(4)             │               │               │
 │               │──────────────▶│               │               │               │
 │               │               │ ButtonClicked(4)──────────────▶               │
 │               │               │               │               │ Save()        │
 │               │               │               │               │──────────────▶│
 │               │               │               │               │ AddAsync()    │
 │               │               │               │               │◀──────────────│
 │               │               │ PreviousControl()◀────────────│               │
 │               │               │ DataLoad(null)▶│               │               │
 │               │               │               │ GetAllAsync()─────────────────▶
 │               │               │               │◀──────────────────────────────│
 │               │               │               │ rebind grid   │               │
 │ [grid refreshed]              │               │               │               │"""
)

# ══════════════════════════════════════════════════════════════════════════════
# 9. PLUGIN / MODULE ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
h1("9. Plugin / Module Architecture for Future Modules")

h2("Core Interface")
code_block(
"""public interface IMedicalModule
{
    string ModuleId              { get; }   // stable, never changes
    string DisplayName           { get; }
    Version Version              { get; }
    string[] RequiredPermissions { get; }   // future role-based access
    IReadOnlyList<MedicalControl> CreateControls(ControlManager manager);
    void RegisterRepositories(IServiceCollection services);  // future DI
}"""
)

h2("Self-Registering Module (Zero Changes to Existing Code)")
code_block(
"""// New "Allergies" module — no modifications to MainForm, ControlManager, or any enum
public class AllergiesModule : IMedicalModule
{
    public string ModuleId    => "allergies";
    public string DisplayName => "Allergies";
    public Version Version    => new(1, 0, 0);
    public string[] RequiredPermissions => [];

    public IReadOnlyList<MedicalControl> CreateControls(ControlManager m) =>
        [new AllergiesControl(m), new AllergyDetailControl(m)];

    public void RegisterRepositories(IServiceCollection s) =>
        s.AddSingleton<IRepository<Allergy>, SqlAllergyRepository>();
}

// Startup registration — only line that ever needs adding
static void Main()
{
    ModuleRegistry registry = new();
    registry.Register(new ImmunizationModule());
    registry.Register(new DoctorVisitModule());
    registry.Register(new AllergiesModule());   // ← just add this
    Application.Run(new LoginForm(registry));
}"""
)

h2("Future: Assembly-Based Plugin Loading")
code_block(
"""public class PluginLoader
{
    public IEnumerable<IMedicalModule> LoadFrom(string pluginDirectory)
    {
        foreach (var dll in Directory.GetFiles(pluginDirectory, "*.Module.dll"))
        {
            var asm = Assembly.LoadFrom(dll);
            foreach (var type in asm.GetTypes()
                .Where(t => typeof(IMedicalModule).IsAssignableFrom(t)
                         && !t.IsAbstract))
            {
                yield return (IMedicalModule)Activator.CreateInstance(type)!;
            }
        }
    }
}"""
)

h2("Module Versioning Strategy")
bullet("Each IMedicalModule declares a Version — the host can enforce minimum version requirements")
bullet("Database migrations tagged per module: Migrations/immunization/v1_0_0.sql")
bullet("Breaking schema changes require a new ModuleId suffix: \"immunization-v2\"")
bullet("Module metadata stored in DB: tblModules (ModuleId, Version, InstalledAt)")

# ══════════════════════════════════════════════════════════════════════════════
# 10. UNIT TESTING STRATEGY
# ══════════════════════════════════════════════════════════════════════════════
h1("10. Unit Testing Strategy")

h2("Test Project Structure")
code_block(
"""prgPMR.Tests/
├── Core/
│   ├── ControlManagerTests.cs
│   └── ModuleRegistryTests.cs
├── Domain/
│   ├── ImmunizationTests.cs
│   └── DoctorVisitTests.cs
├── Modules/
│   ├── ImmunizationControlTests.cs
│   └── ImmunizationDetailControlTests.cs
├── Data/
│   ├── SqlImmunizationRepositoryTests.cs   (integration — needs DB)
│   └── InMemoryRepositoryTests.cs          (pure unit)
└── Helpers/
    ├── FakeControlManager.cs
    └── FakeRepository.cs"""
)

h2("Key Test Cases")
code_block(
"""// ── ControlManager ────────────────────────────────────────────
[Fact]
public void ButtonClicked_RoutesToActiveControlAction()
{
    bool called = false;
    var manager = BuildManagerWithAction(() => called = true);
    manager.ButtonClicked(0);
    Assert.True(called);
}

[Fact]
public void NextControl_AtLastControl_DoesNotAdvance()
{
    var manager = BuildManagerWithOneControl();
    manager.NextControl(null);
    Assert.Equal(0, manager.ActiveControlIndex);
}

[Fact]
public void SetVisible_False_HidesAllControls()
{
    var manager = BuildImmunizationManager();
    manager.SetVisible(false);
    Assert.All(manager.MedicalControls, c => Assert.False(c.Visible));
}

// ── ModuleRegistry ────────────────────────────────────────────
[Fact]
public void Register_DuplicateModuleId_Throws()
{
    var registry = new ModuleRegistry();
    registry.Register(new ImmunizationModule());
    Assert.Throws<InvalidOperationException>(
        () => registry.Register(new ImmunizationModule()));
}

// ── InMemory Repository ───────────────────────────────────────
[Fact]
public async Task AddAsync_ThenGetAll_ReturnsRecord()
{
    var repo      = new InMemoryImmunizationRepository();
    var patientId = Guid.NewGuid();
    var record    = new Immunization { PatientId = patientId };
    await repo.AddAsync(record);
    var results   = await repo.GetAllAsync(patientId);
    Assert.Single(results);
    Assert.Equal(record.Id, results[0].Id);
}

[Fact]
public async Task DeleteAsync_NonExistentId_Throws()
{
    var repo = new InMemoryImmunizationRepository();
    await Assert.ThrowsAsync<KeyNotFoundException>(
        () => repo.DeleteAsync(Guid.NewGuid()));
}

// ── Bug regression: Cancel on clean form ─────────────────────
[Fact]
public void Cancel_WhenFormIsClean_NavigatesBack()
{
    // Currently FAILS — documents the existing bug
    var fakeManager = new FakeControlManager();
    var control     = new ImmunizationDetailControl(fakeManager, new FakeRepository());
    control.DataLoad(new ImmunizationPayload(Mode: ControlPayloadMode.Add));
    // Do NOT modify any field — form is clean
    control.Cancel();
    Assert.True(fakeManager.PreviousControlCalled);
}"""
)

h2("Testing Boundaries")
make_table(
    ["Layer", "Test Type", "Notes"],
    [
        ["Domain objects",         "Unit",       "Pure C# — no WinForms dependency"],
        ["ControlManager",         "Unit",       "Use FakeControlManager and FakeModule"],
        ["ModuleRegistry",         "Unit",       "Simple dictionary contract"],
        ["InMemory repositories",  "Unit",       "Fast, no infrastructure"],
        ["SQL repositories",       "Integration","Requires SQL Server or TestContainers"],
        ["Designer code",          "None",       "Do not unit-test generated layout code"],
        ["Visual appearance",      "Manual/UI",  "Use FlaUI or WinAppDriver for UI tests"],
    ],
    col_widths=[1.8, 1.3, 3.3]
)

# ══════════════════════════════════════════════════════════════════════════════
# 11. DEVELOPMENT ROADMAP
# ══════════════════════════════════════════════════════════════════════════════
h1("11. Development Roadmap")

h2("Phase 1 — Foundation Fixes (Do Now, Low Risk)")
make_table(
    ["Priority", "Task"],
    [
        ["P0", "Fix pgmPMRMain.cs to launch LoginForm instead of MainForm"],
        ["P0", "Fix connection string typo ('GridInitial Catalog' → 'Initial Catalog')"],
        ["P0", "Move connection string to appsettings.json or user settings"],
        ["P1", "Fix Cancel() bug — always navigate back regardless of dirty state"],
        ["P1", "Fix Delete() — move button reset to after user confirms"],
        ["P1", "Rename DoctorVisitsDetail → DoctorVisitsDetailControl"],
        ["P1", "Remove debug artifacts (label text responses in TestsControl)"],
        ["P2", "Remove empty auto-generated event handlers"],
    ],
    col_widths=[0.8, 5.6]
)

h2("Phase 2 — Architecture Hardening (Next Sprint)")
make_table(
    ["Priority", "Task"],
    [
        ["P1", "Extract MedicalControlType enum out of MainForm; move to ModuleRegistry"],
        ["P1", "Create IMedicalModule interface; refactor ControlManager switch to use it"],
        ["P1", "Create IRepository<T> + InMemoryXxxRepository stubs for all modules"],
        ["P1", "Wire DoctorVisitsControl and DoctorVisitsDetailControl to match Immunization"],
        ["P2", "Create test project; write ControlManager and ModuleRegistry tests"],
        ["P2", "Introduce SqlConnectionFactory; inject rather than hardcode"],
    ],
    col_widths=[0.8, 5.6]
)

h2("Phase 3 — Domain Model (Next 2–4 Weeks)")
make_table(
    ["Priority", "Task"],
    [
        ["P1", "Define all domain classes (Immunization, DoctorVisit, Medication, etc.)"],
        ["P1", "Replace ImmunizationAddData(SampleText) with real ImmunizationPayload"],
        ["P1", "Wire ImmunizationDetailControl.Save() to real SqlImmunizationRepository"],
        ["P2", "Wire ImmunizationControl.InitializeGrid() to repository (remove DataTable literal)"],
        ["P2", "Implement PatientProfile; replace hardcoded user text fields in MainForm"],
    ],
    col_widths=[0.8, 5.6]
)

h2("Phase 4 — Remaining Modules")
make_table(
    ["Priority", "Task"],
    [
        ["P1", "Implement DoctorVisitsControl / DoctorVisitsDetailControl fully"],
        ["P2", "Implement MedicationsControl"],
        ["P2", "Implement HospitalizationsControl + detail"],
        ["P3", "Implement SurgeriesControl, BloodworkControl, TestsControl, FamilyHistoryControl"],
    ],
    col_widths=[0.8, 5.6]
)

h2("Phase 5 — Security & Polish")
make_table(
    ["Priority", "Task"],
    [
        ["P0", "Hash passwords in DB (bcrypt or PBKDF2) — never store plaintext"],
        ["P1", "Use ExecuteScalar, not SqlDataAdapter.Fill, for login query"],
        ["P1", "Pass authenticated user context (PatientProfile) from LoginForm to MainForm"],
        ["P2", "Attachment support (PDF/image upload) in detail views"],
        ["P2", "Export records to PDF or CSV"],
        ["P3", "Full plugin loader from external assemblies"],
        ["P3", "Automated UI tests (FlaUI or WinAppDriver)"],
    ],
    col_widths=[0.8, 5.6]
)

# ══════════════════════════════════════════════════════════════════════════════
# SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h1("Summary — Biggest Wins")

make_table(
    ["Action", "Impact"],
    [
        ["Fix pgmPMRMain to launch LoginForm",          "Authentication actually works"],
        ["Fix connection string typo",                  "Application can connect to the database"],
        ["Extract IMedicalModule + ModuleRegistry",     "New modules in 1 file, 1 line of registration"],
        ["Add IRepository<T> + in-memory stubs",        "All controls testable without a live database"],
        ["Fix Cancel() bug",                            "Core UX flow no longer broken on clean form"],
        ["Move enum out of MainForm",                   "Shell decoupled from module identity"],
        ["Hash passwords",                              "Critical security baseline met"],
    ],
    col_widths=[3.0, 3.4]
)

doc.add_paragraph()
body(
    "The core navigation pattern (ControlManager → MedicalControl → DataInterface dispatch) is "
    "a genuinely good design — it is the skeleton worth preserving and building on. "
    "The main work is carving out a clean data layer, replacing the hardcoded module switch "
    "with a self-registering module registry, fixing the three critical bugs (bypassed login, "
    "malformed connection string, plaintext passwords), and then adding the remaining module "
    "implementations on top of that solid frame.",
    italic=True
)

# ── Save ─────────────────────────────────────────────────────────────────────
output_path = "E:/Dev/GitHub/medical-dashboard/PMR_Architecture_Review.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
