"""
Generates the ControlManager Redesign Word document.
Run: python generate_controlmanager_redesign.py
Output: ControlManager_Redesign.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.1)
    section.right_margin  = Inches(1.1)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY    = RGBColor(0x1F, 0x39, 0x64)
BLUE    = RGBColor(0x2E, 0x74, 0xB5)
TEAL    = RGBColor(0x00, 0x6B, 0x6B)
RED     = RGBColor(0xC0, 0x00, 0x00)
ORANGE  = RGBColor(0xBF, 0x59, 0x00)
GREEN   = RGBColor(0x37, 0x86, 0x30)
PURPLE  = RGBColor(0x5B, 0x2C, 0x8D)
GRAY    = RGBColor(0x59, 0x59, 0x59)
LGRAY   = RGBColor(0x99, 0x99, 0x99)
WHITE   = RGBColor(0xFF, 0xFF, 0xFF)

HEX_LIGHT_RED    = "FCEAEA"
HEX_LIGHT_GREEN  = "EAF4EA"
HEX_LIGHT_BLUE   = "EBF3FB"
HEX_LIGHT_PURPLE = "F3EEF9"
HEX_LIGHT_GRAY   = "F5F5F5"
HEX_AMBER        = "FFF3CD"

# ── Cell background ───────────────────────────────────────────────────────────
def cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

# ── Heading helpers ───────────────────────────────────────────────────────────
def h1(text, color=NAVY):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = color
    p.runs[0].font.size = Pt(17)
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(6)

def h2(text, color=BLUE):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = color
    p.runs[0].font.size = Pt(13)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)

def h3(text, color=BLUE):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.color.rgb = color
    p.runs[0].font.size = Pt(11)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)

def body(text, bold=False, italic=False, color=None, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    run = p.add_run(text)
    run.font.size  = Pt(10.5)
    run.bold       = bold
    run.italic     = italic
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(4)
    return p

def mixed(parts):
    """Parts: list of (text, bold, italic, color). Returns paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    for (text, bold, italic, color) in parts:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.size = Pt(10.5)
        if color:
            r.font.color.rgb = color
    return p

def bullet(text, level=0, color=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(2)
    return p

def numbered(text, level=0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(2)
    return p

def callout(label, text, bg_hex, label_color):
    """A coloured callout box."""
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell = t.rows[0].cells[0]
    cell_bg(cell, bg_hex)
    p = cell.paragraphs[0]
    r1 = p.add_run(f"{label}  ")
    r1.bold = True
    r1.font.color.rgb = label_color
    r1.font.size = Pt(10.5)
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()

def insight(text):
    callout("💡 Key Insight:", text, HEX_LIGHT_BLUE, BLUE)

def warning(text):
    callout("⚠ Problem:", text, HEX_LIGHT_RED, RED)

def note(text):
    callout("📌 Note:", text, HEX_AMBER, ORANGE)

def code_block(text, title=None):
    if title:
        tp = doc.add_paragraph()
        tp.paragraph_format.left_indent = Inches(0.25)
        tr = tp.add_run(f"  {title}")
        tr.bold = True
        tr.font.size = Pt(9)
        tr.font.color.rgb = GRAY
        tp.paragraph_format.space_after = Pt(0)
        tp.paragraph_format.space_before = Pt(4)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(8)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F0F0F0")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name      = "Courier New"
    run.font.size      = Pt(8.5)
    run.font.color.rgb = TEAL
    return p

def side_by_side(left_title, left_code, right_title, right_code,
                 left_bg="FCEAEA", right_bg="EAF4EA",
                 left_hdr="1F3964", right_hdr="375E1E"):
    t = doc.add_table(rows=2, cols=2)
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT

    for ci, (title, hdr) in enumerate([(left_title, left_hdr), (right_title, right_hdr)]):
        c = t.rows[0].cells[ci]
        cell_bg(c, hdr)
        r = c.paragraphs[0].add_run(title)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(9.5)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for ci, (code, bg) in enumerate([(left_code, left_bg), (right_code, right_bg)]):
        c = t.rows[1].cells[ci]
        cell_bg(c, bg)
        p = c.paragraphs[0]
        p.paragraph_format.left_indent = Inches(0.05)
        run = p.add_run(code)
        run.font.name      = "Courier New"
        run.font.size      = Pt(7.8)
        run.font.color.rgb = TEAL

    for row in t.rows:
        row.cells[0].width = Inches(3.05)
        row.cells[1].width = Inches(3.05)

    doc.add_paragraph()

def make_table(headers, rows, col_widths=None, row_colors=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr_row = t.rows[0]
    for i, h in enumerate(headers):
        c = hdr_row.cells[i]
        cell_bg(c, "1F3964")
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.color.rgb = WHITE
        r.font.size = Pt(10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        bg = row_colors[ri] if row_colors else (HEX_LIGHT_BLUE if ri % 2 == 1 else None)
        for ci, val in enumerate(row):
            c = tr.cells[ci]
            if bg:
                cell_bg(c, bg)
            r = c.paragraphs[0].add_run(str(val))
            r.font.size = Pt(9.5)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def component_header(number, name, role, color_hex="1F3964"):
    t = doc.add_table(rows=1, cols=1)
    t.style = "Table Grid"
    cell_bg(t.rows[0].cells[0], color_hex)
    p = t.rows[0].cells[0].paragraphs[0]
    r1 = p.add_run(f"Component {number}:  ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(0xAA, 0xCC, 0xFF)
    r1.font.size = Pt(10)
    r2 = p.add_run(name)
    r2.bold = True
    r2.font.color.rgb = WHITE
    r2.font.size = Pt(13)
    r3 = p.add_run(f"   —   {role}")
    r3.italic = True
    r3.font.color.rgb = RGBColor(0xCC, 0xDD, 0xFF)
    r3.font.size = Pt(10)
    doc.add_paragraph()

def divider():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(50)
r = p.add_run("ControlManager")
r.bold = True; r.font.size = Pt(30); r.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("A Ground-Up Redesign")
r2.font.size = Pt(18); r2.font.color.rgb = BLUE

doc.add_paragraph()

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
for line in [
    "Project:  Personal Medical Records Manager (prgPMR)",
    "File:     prgPMR/ControlManager.cs",
    "Review Date:  February 2026",
]:
    r3 = p3.add_run(line + "\n")
    r3.font.size = Pt(10.5); r3.font.color.rgb = GRAY; r3.font.name = "Courier New"

doc.add_paragraph()
p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
r4 = p4.add_run(
    "Not an incremental fix — a clean-room redesign that identifies what "
    "ControlManager is really trying to be,\nthen builds it correctly from first principles."
)
r4.italic = True; r4.font.size = Pt(10.5); r4.font.color.rgb = GRAY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1: REFACTOR VS REDESIGN
# ══════════════════════════════════════════════════════════════════════════════
h1("1. Refactor vs Redesign — Why This Needs a Clean Sheet")

body(
    "The previous document listed 7 targeted fixes to ControlManager. Those fixes are correct "
    "and should be made. But they are plasters — they leave the fundamental structural problems "
    "in place. This document asks a different question:"
)
body(
    "If we started with a blank file and only the requirements, what would we write?",
    bold=True, italic=True, color=BLUE
)

make_table(
    ["Dimension", "Refactor Approach", "Redesign Approach"],
    [
        ["Starting point",     "Existing code, fix what is wrong",    "Requirements only, no legacy assumptions"],
        ["Switch statement",   "Replace with IMedicalModule",          "Never existed — modules are plugins from day one"],
        ["Navigation",         "Index becomes a stack",                "A dedicated ModuleNavigator state machine"],
        ["Button management",  "Tighten the parallel arrays",          "ButtonDescriptor eliminates parallel arrays entirely"],
        ["Circular dependency","MedicalControl still calls back in",   "Broken completely — views raise events, nothing calls back"],
        ["MedicalControl",     "Keeps its current shape, improved",    "Evolves to implement IModuleView — cleaner contract"],
        ["Testability",        "Add a read-only property",             "Every component is independently testable in isolation"],
        ["MainForm coupling",  "Reduced but still present",            "ModuleCoordinator absorbs all routing from MainForm"],
    ],
    col_widths=[1.5, 2.3, 2.3]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2: WHAT IS CONTROLMANAGER REALLY DOING?
# ══════════════════════════════════════════════════════════════════════════════
h1("2. What Is ControlManager Really Doing?")

body(
    "Before designing the replacement, we need to name what ControlManager actually is. "
    "Read the current code and count its responsibilities:"
)

make_table(
    ["Responsibility", "Evidence in Code", "Should Belong To"],
    [
        ["Factory — creates concrete controls",      "switch(type) { case Immunization: new ImmunizationControl... }",  "IMedicalModule"],
        ["Navigation state — tracks active view",    "private int activeControl",                                        "ModuleNavigator"],
        ["View lifecycle — activates / deactivates", "MedicalControls[activeControl].DataLoad(data)",                    "ModuleNavigator"],
        ["Visibility — shows / hides panels",        "foreach(m) m.Visible = ...",                                       "ModuleHost"],
        ["Button bar sync — updates shared buttons", "Buttons[i].Text = m.ButtonsText[i]",                               "ButtonBarController"],
        ["Button routing — dispatches click events", "MedicalControls[activeControl].ButtonActions[index]()",            "ButtonBarController"],
        ["Module identity — knows all module types", "MainForm.MedicalControlType dependency",                           "ModuleRegistry"],
    ],
    col_widths=[1.9, 2.5, 1.7]
)

insight(
    "ControlManager is doing seven distinct jobs. The redesign gives each job to a "
    "dedicated component. ControlManager itself becomes a thin 5-line coordinator — "
    "ModuleHost — that just wires the others together."
)

body(
    "The closest analogy in modern frameworks is a navigation controller:",
    italic=True
)
bullet("iOS: UINavigationController + UINavigationBar")
bullet("WPF: Frame + NavigationService")
bullet("Xamarin/MAUI: NavigationPage + NavigationStack")
bullet("Web: Router + active route component")
body(
    "All of these are a navigation state machine (push/pop views) + a UI element that reflects "
    "the active view's toolbar state. That is exactly what ControlManager should be.",
    italic=True
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3: THE PARALLEL ARRAYS PROBLEM
# ══════════════════════════════════════════════════════════════════════════════
h1("3. The Parallel Arrays Problem (Fixing the Root Before the Branches)")

body(
    "Before describing the new components, there is a foundational issue in MedicalControl "
    "that the redesign must solve first, because everything else builds on top of it."
)

warning(
    "ButtonsText (string[]) and ButtonActions (Action[]) are parallel arrays. "
    "They must always be the same length, and null positions must align between them. "
    "The only guard is a Debug.Assert — which is silent in Release builds. "
    "One misalignment causes a silent wrong button label or an unhandled exception."
)

side_by_side(
    "Current — Parallel Arrays  ✗",
"""// Two separate arrays that must stay in sync
private string[] _buttonText    = [];
private Action[] _buttonActions = [];

// Both must have matching nulls at same positions
SetButtons(
  ["Add", "Edit", null, "Reset", null, null],
  [ Add,   Edit,  null,  Reset,  null,  null]
);
// Miss-align these and you get:
// - wrong label on a button
// - NullReferenceException on click
// - no compile-time or runtime warning""",
    "Redesign — ButtonDescriptor  ✓",
"""// One record that holds label + action together
public record ButtonDescriptor(
    string? Label,
    Action? Action,
    bool    IsEnabled = true)
{
    public static readonly ButtonDescriptor Empty
        = new(null, null);
    public bool IsVisible => Label is not null;
}

// Cannot get out of sync — they are the same object
SetButtons([
    new("Add",    Add),
    new("Edit",   Edit),
    ButtonDescriptor.Empty,
    new("Reset",  Reset),
    ButtonDescriptor.Empty,
    ButtonDescriptor.Empty,
]);"""
)

body("This is the single change with the most downstream impact — it simplifies "
     "MedicalControl, ButtonBarController, and every concrete control class.", italic=True)

h2("ButtonPresets — Static Factory for Common Configurations")
body(
    "The current design has each concrete control hold its own "
    "lowerbuttonBarPresetTextsDict and lowerbuttonBarPresetActionDict. "
    "These are duplicated in ImmunizationControl, ImmunizationDetailControl, and every "
    "future module. The redesign consolidates them into a single static class."
)

side_by_side(
    "Current — Duplicated in every control  ✗",
"""// In ImmunizationControl:
lowerbuttonBarPresetActionDict =
  new Dictionary<
    LowerbuttonBarPresetGrouping,
    Action[]>
  {
    { GridInitial,
      [Add, null, null, Reset,
       null, null] },
    { GridSelect,
      [Add, Edit, Delete, Reset,
       null, null] },
  };

// Then to use it:
SetButtons(
  lowerbuttonBarPresetTextsDict[GridInitial],
  lowerbuttonBarPresetActionDict[GridInitial]
);

// ImmunizationDetailControl has its own copy
// Every future control will have its own copy""",
    "Redesign — Single static class  ✓",
"""public static class ButtonPresets
{
    public static IReadOnlyList<ButtonDescriptor>
        GridInitial(Action add, Action reset) =>
    [
        Btn("Add",   add),   Empty, Empty,
        Btn("Reset", reset), Empty, Empty
    ];

    public static IReadOnlyList<ButtonDescriptor>
        GridSelect(Action add, Action edit,
                   Action delete, Action reset) =>
    [
        Btn("Add",    add),
        Btn("Edit",   edit),
        Btn("Delete", delete),
        Btn("Reset",  reset),
        Empty, Empty
    ];

    public static IReadOnlyList<ButtonDescriptor>
        DetailAdd(Action reset, Action save,
                  Action cancel) => [ ... ];

    public static IReadOnlyList<ButtonDescriptor>
        DetailEdit(Action delete, Action reset,
                   Action save, Action cancel)
            => [ ... ];
}"""
)

body("Usage in a concrete control drops from 12 lines to 1:", italic=True)
code_block(
"""// Old (in ImmunizationControl):
SetButtons(
    lowerbuttonBarPresetTextsDict[LowerbuttonBarPresetGrouping.GridInitial],
    lowerbuttonBarPresetActionDict[LowerbuttonBarPresetGrouping.GridInitial]);

// New:
SetButtons(ButtonPresets.GridInitial(Add, Reset));""",
    "Before vs After — using a preset"
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4: THE FIVE COMPONENTS
# ══════════════════════════════════════════════════════════════════════════════
h1("4. The Five Components — Overview")

body(
    "The redesign splits ControlManager's seven responsibilities across five purpose-built "
    "components. Each component has exactly one job and no knowledge of the others "
    "except through well-defined interfaces and events."
)

code_block(
"""┌─────────────────────────────────────────────────────────────────────────────┐
│                            MainForm  (shell)                                 │
│  nav panel buttons  │  pnlMain (panel host)  │  shared bottom button bar    │
└──────────┬──────────┴────────────┬────────────┴──────────────┬───────────────┘
           │ activates             │ hosts controls             │ click events
           ▼                       │                            ▼
┌──────────────────────┐           │              ┌─────────────────────────────┐
│   ModuleCoordinator  │           │              │     ButtonBarController      │
│  - activates modules │           │              │  - subscribes to Navigator   │
│  - routes btn clicks │           │              │  - syncs Button[] text/vis   │
│  - owns the registry │           │              │  - routes Action on click    │
└──────────┬───────────┘           │              └─────────────────────────────┘
           │ creates                │                           ▲ subscribes
           ▼                       │                           │
┌──────────────────────┐           │              ┌─────────────────────────────┐
│      ModuleHost      │           │              │       ModuleNavigator        │
│  (new ControlManager)│           │              │  - Stack<IModuleView>        │
│  - owns Navigator    │◀──────────┘              │  - Push / Pop / Reset        │
│  - owns module views │  GetControls()           │  - fires Navigated event     │
│  - manages visibility│                          └─────────────────────────────┘
└──────────────────────┘                                       ▲ Push/Pop called by
                                                               │
                              ┌────────────────────────────────┘
                              │
              ┌───────────────┴──────────────────────────────┐
              │           MedicalControl  (implements IModuleView)            │
              │  - UserControl subclass (WinForms Designer safe)              │
              │  - Raises ButtonsChanged event (no more Manager callback)     │
              │  - OnActivated / OnReactivated / OnDeactivated lifecycle      │
              │  - Uses ButtonPresets for clean button configuration          │
              └──────────────────────────────────────────────┘""",
    "Component Relationship Diagram"
)

make_table(
    ["Component", "Single Responsibility", "Replaces"],
    [
        ["ButtonDescriptor",    "Pair a label with its action in one cohesive unit",             "Parallel string[] + Action[] arrays"],
        ["ButtonPresets",       "Named factories for common button bar configurations",          "lowerbuttonBarPresetTextsDict + ActionDict in every control"],
        ["ModuleNavigator",     "Pure navigation state machine (push/pop view stack)",           "int activeControl + NextControl/PreviousControl"],
        ["ButtonBarController", "Observe navigator + active view, sync physical buttons",        "SetVisible button-update loop + ButtonClicked routing"],
        ["ModuleHost",          "Thin coordinator: own navigator + views + module visibility",   "The bulk of ControlManager"],
        ["ModuleCoordinator",   "Activate/deactivate modules, route shell events",              "Module management logic in MainForm"],
    ],
    col_widths=[1.65, 2.65, 2.3]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5: COMPONENT 1 — IModuleView + MedicalControl
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
component_header(1, "IModuleView + MedicalControl",
                 "Contract for all navigable views — replaces the Manager back-reference",
                 "1F3964")

h2("The Circular Dependency Problem")
warning(
    "In the current design, MedicalControl holds a reference to ControlManager "
    "(Manager property) and calls Manager.RefreshVisibility() every time buttons change. "
    "This means: ControlManager owns MedicalControl, AND MedicalControl calls back into "
    "ControlManager. Both sides depend on each other — a circular dependency that makes "
    "either class impossible to test in isolation."
)

code_block(
"""// Current circular dependency:
ControlManager owns ──────────────────────────────▶ MedicalControl
MedicalControl.Manager calls back into ◀─────────── ControlManager

// Cannot test MedicalControl without ControlManager
// Cannot test ControlManager without MedicalControl""",
    "The Circular Dependency"
)

h2("The Fix: Events Break the Circle")
body(
    "Instead of MedicalControl calling back into its owner, it raises a ButtonsChanged event. "
    "ButtonBarController (not ControlManager) subscribes to this event. "
    "MedicalControl no longer knows who is listening — the dependency is inverted."
)

code_block(
"""// New: event-driven, no back-reference
MedicalControl raises ───────────────────────────▶ ButtonsChanged event
ButtonBarController subscribes to ◀────────────── ButtonsChanged event

// MedicalControl has zero knowledge of who is listening
// Both are independently testable""",
    "Circular Dependency Broken"
)

h2("IModuleView Interface")
code_block(
"""namespace prgPMR.Core
{
    /// <summary>
    /// Contract for any navigable view within a medical module.
    /// Implemented by all MedicalControl subclasses.
    /// </summary>
    public interface IModuleView
    {
        // ── Identity ─────────────────────────────────────────────
        /// <summary>Stable ID used for navigation lookups.</summary>
        string ViewId { get; }

        /// <summary>The underlying WinForms control to host in pnlMain.</summary>
        UserControl Control { get; }

        // ── Button state (observable) ─────────────────────────────
        /// <summary>Current button bar configuration for this view.</summary>
        IReadOnlyList<ButtonDescriptor> ButtonDescriptors { get; }

        /// <summary>
        /// Fired whenever ButtonDescriptors changes.
        /// ButtonBarController subscribes to keep the physical bar in sync.
        /// </summary>
        event EventHandler? ButtonsChanged;

        // ── Lifecycle ─────────────────────────────────────────────
        /// <summary>Called when this view is pushed onto the stack (forward nav).</summary>
        void OnActivated(IControlPayload? payload);

        /// <summary>Called when a child view is popped and we return here.</summary>
        void OnReactivated();

        /// <summary>Called when navigating away from this view.</summary>
        void OnDeactivated();
    }
}"""
)

h2("Updated MedicalControl (implements IModuleView)")
note(
    "MedicalControl still extends UserControl — the WinForms Designer attribute is preserved. "
    "The only changes are: (1) it implements IModuleView, (2) SetButtons takes ButtonDescriptor "
    "instead of parallel arrays, (3) the Manager back-reference is gone."
)
code_block(
"""namespace prgPMR.Core
{
    [TypeDescriptionProvider(
        typeof(AbstractControlDescriptionProvider<MedicalControl, UserControl>))]
    public abstract class MedicalControl : UserControl, IModuleView
    {
        // ── IModuleView: Identity ─────────────────────────────────
        public abstract string ViewId { get; }
        UserControl IModuleView.Control => this;   // 'this' IS the control

        // ── IModuleView: Button state ─────────────────────────────
        private IReadOnlyList<ButtonDescriptor> _descriptors = [];
        public  IReadOnlyList<ButtonDescriptor> ButtonDescriptors => _descriptors;
        public  event EventHandler? ButtonsChanged;

        // ── IModuleView: Lifecycle ────────────────────────────────
        public virtual void OnActivated(IControlPayload? payload) { }
        public virtual void OnReactivated() { }
        public virtual void OnDeactivated() { }

        // ── Navigation (injected, not back-referenced) ─────────────
        // Views call navigator.Push/Pop instead of Manager.Next/Prev
        protected IModuleNavigator Navigator { get; }

        protected MedicalControl(IModuleNavigator navigator)
        {
            Navigator = navigator;
        }

        // ── Button management ─────────────────────────────────────
        /// <summary>
        /// Set the button bar configuration for this view.
        /// Raises ButtonsChanged — ButtonBarController responds automatically.
        /// No more Manager.RefreshVisibility() call.
        /// </summary>
        protected void SetButtons(IReadOnlyList<ButtonDescriptor> descriptors)
        {
            _descriptors = descriptors;
            ButtonsChanged?.Invoke(this, EventArgs.Empty);
        }
    }
}"""
)

h2("How a Concrete Control Changes")
side_by_side(
    "ImmunizationControl — Before  ✗",
"""public partial class ImmunizationControl
    : MedicalControl
{
    // Duplicate preset dictionaries
    private Dictionary<
      LowerbuttonBarPresetGrouping,
      Action[]>
        lowerbuttonBarPresetActionDict;

    public ImmunizationControl(
        ControlManager m) : base(m)
    {
        InitializeComponent();
        lowerbuttonBarPresetActionDict =
            new() {
              { GridInitial,
                [Add, null, null,
                 Reset, null, null] },
            };
        InitializeGrid();
    }

    public override void DataLoad(
        DataInterface? data)
    {
        if (data == null)
            InitializeGrid();  // null = reset
        else
            throw new ArgumentException(
                "Bad Immunization Data");
    }

    public void Add() =>
        Manager.NextControl(
            new ImmunizationAddData("test"));
}""",
    "ImmunizationControl — After  ✓",
"""public partial class ImmunizationControl
    : MedicalControl
{
    // ViewId from IModuleView
    public override string ViewId =>
        "immunization.list";

    public ImmunizationControl(
        IModuleNavigator navigator)
        : base(navigator)
    {
        InitializeComponent();
        InitializeGrid();
    }

    // DataLoad replaced by lifecycle methods
    public override void OnActivated(
        IControlPayload? _)
    {
        // payload is null for list view
        InitializeGrid();
    }

    public override void OnReactivated()
    {
        // Returning from detail — refresh list
        InitializeGrid();
    }

    private void InitializeGrid()
    {
        SetButtons(ButtonPresets
            .GridInitial(Add, Reset));
        // ... load data
    }

    private void Add() =>
        Navigator.Push(
            _detailView,   // injected ref
            new ImmunizationPayload(
                ControlPayloadMode.Add));
}"""
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6: COMPONENT 2 — ModuleNavigator
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
component_header(2, "ModuleNavigator",
                 "Pure navigation state machine — a stack of IModuleView",
                 "1F3964")

h2("Why a Stack, Not an Index?")

make_table(
    ["", "Index (current)", "Stack (redesign)"],
    [
        ["Representation",   "int activeControl (position in a List)",     "Stack<IModuleView>"],
        ["Navigate forward", "activeControl++ — breaks if list order changes", "Push(view, payload) — always correct"],
        ["Navigate back",    "activeControl-- — breaks if at 0",           "Pop() — naturally bounded, always safe"],
        ["Multi-level nav",  "Not possible — only 2 levels (list + detail)", "Any depth — detail could push a sub-detail"],
        ["State clarity",    "Must infer meaning from index value",        "Stack top IS the current view — self-evident"],
        ["Reset to root",    "activeControl = 0 (magic number)",           "while(CanGoBack) Pop() -- explicit"],
        ["Testability",      "Cannot observe state without reflection",    "Depth, Current, CanGoBack are read-only properties"],
    ],
    col_widths=[1.45, 2.25, 2.4]
)

h2("IModuleNavigator Interface")
code_block(
"""namespace prgPMR.Core
{
    /// <summary>
    /// Exposes the navigation API to IModuleView implementations.
    /// Views depend on this interface, not on ModuleNavigator directly.
    /// This keeps views testable with a fake navigator.
    /// </summary>
    public interface IModuleNavigator
    {
        IModuleView? Current { get; }
        bool CanGoBack       { get; }
        int  Depth           { get; }

        void Push(IModuleView view, IControlPayload? payload);
        void Pop();
        void Reset();

        event EventHandler<NavigatedEventArgs>? Navigated;
    }

    public class NavigatedEventArgs : EventArgs
    {
        public IModuleView?          Previous  { get; }
        public IModuleView?          Current   { get; }
        public NavigationDirection   Direction { get; }

        public NavigatedEventArgs(
            IModuleView? previous,
            IModuleView? current,
            NavigationDirection direction)
        {
            Previous  = previous;
            Current   = current;
            Direction = direction;
        }
    }

    public enum NavigationDirection { Forward, Back, Reset }
}"""
)

h2("ModuleNavigator Implementation")
code_block(
"""namespace prgPMR.Core
{
    public class ModuleNavigator : IModuleNavigator
    {
        private readonly Stack<IModuleView> _stack = new();

        // ── State (read-only) ────────────────────────────────────
        public IModuleView? Current   => _stack.TryPeek(out var v) ? v : null;
        public bool         CanGoBack => _stack.Count > 1;
        public int          Depth     => _stack.Count;

        // ── Events ───────────────────────────────────────────────
        public event EventHandler<NavigatedEventArgs>? Navigated;

        // ── Navigation ───────────────────────────────────────────

        /// <summary>
        /// Navigate forward: deactivate current, push new view, activate it.
        /// </summary>
        public void Push(IModuleView view, IControlPayload? payload)
        {
            var previous = Current;
            previous?.OnDeactivated();     // tell old view it is leaving

            _stack.Push(view);
            view.OnActivated(payload);     // tell new view what it should show

            Navigated?.Invoke(this, new NavigatedEventArgs(
                previous, view, NavigationDirection.Forward));
        }

        /// <summary>
        /// Navigate back: deactivate current, pop stack, reactivate previous.
        /// </summary>
        public void Pop()
        {
            if (!CanGoBack) return;

            var leaving = _stack.Pop();
            leaving.OnDeactivated();       // tell detail view it is closing

            Current?.OnReactivated();      // tell list view to refresh

            Navigated?.Invoke(this, new NavigatedEventArgs(
                leaving, Current, NavigationDirection.Back));
        }

        /// <summary>
        /// Snap back to root view — used when the module loses focus.
        /// </summary>
        public void Reset()
        {
            if (Depth <= 1) return;

            while (_stack.Count > 1)
                _stack.Pop();              // silently unwind — no lifecycle calls

            Current?.OnReactivated();      // reload root view

            Navigated?.Invoke(this, new NavigatedEventArgs(
                null, Current, NavigationDirection.Reset));
        }
    }
}"""
)

note(
    "OnReactivated() is distinct from OnActivated(). When navigating back, the list view "
    "does not receive a payload — it just needs to refresh its data. Having separate methods "
    "makes this intent explicit and avoids the current null-as-reset-signal pattern."
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7: COMPONENT 3 — ButtonBarController
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
component_header(3, "ButtonBarController",
                 "Observer — listens to the navigator and active view, syncs the physical buttons",
                 "1F3964")

h2("Why a Separate Controller?")
body(
    "In the current design, button bar synchronisation is buried inside SetVisible(). "
    "This means navigation logic and button rendering are tangled in the same method. "
    "ButtonBarController extracts button management into a dedicated class that responds "
    "to two event sources:"
)
bullet("ModuleNavigator.Navigated — tells it which view is now active")
bullet("IModuleView.ButtonsChanged — tells it the active view's buttons have changed")
body(
    "This is the Observer pattern. ButtonBarController never polls — it always reacts.",
    italic=True
)

code_block(
"""namespace prgPMR.Core
{
    /// <summary>
    /// Subscribes to a ModuleNavigator and the active IModuleView.
    /// Keeps the physical WinForms Button[] in sync with the active view's
    /// ButtonDescriptors at all times. Also routes button clicks.
    /// </summary>
    public class ButtonBarController
    {
        private readonly Button[]  _buttons;
        private          IModuleView? _activeView;

        public ButtonBarController(Button[] buttons)
        {
            _buttons = buttons;
        }

        // ── Attach / detach from a navigator ─────────────────────
        public void AttachTo(IModuleNavigator navigator)
        {
            navigator.Navigated += OnNavigated;
        }

        public void DetachFrom(IModuleNavigator navigator)
        {
            navigator.Navigated -= OnNavigated;
            DetachFromView(_activeView);
            _activeView = null;
            ClearButtons();             // blank the bar when module is hidden
        }

        // ── Event handlers ────────────────────────────────────────
        private void OnNavigated(object? sender, NavigatedEventArgs e)
        {
            DetachFromView(_activeView);         // stop listening to old view
            _activeView = e.Current;
            AttachToView(_activeView);           // start listening to new view
            SyncButtons();
        }

        private void OnButtonsChanged(object? sender, EventArgs e)
        {
            SyncButtons();                        // active view changed its buttons
        }

        private void AttachToView(IModuleView? view)
        {
            if (view != null)
                view.ButtonsChanged += OnButtonsChanged;
        }

        private void DetachFromView(IModuleView? view)
        {
            if (view != null)
                view.ButtonsChanged -= OnButtonsChanged;
        }

        // ── Button sync ───────────────────────────────────────────
        private void SyncButtons()
        {
            var descriptors = _activeView?.ButtonDescriptors
                              ?? Array.Empty<ButtonDescriptor>();

            for (int i = 0; i < _buttons.Length; i++)
            {
                var d = i < descriptors.Count
                    ? descriptors[i]
                    : ButtonDescriptor.Empty;

                _buttons[i].Visible = d.IsVisible;
                _buttons[i].Enabled = d.IsEnabled;
                _buttons[i].Text    = d.Label ?? string.Empty;
            }
        }

        private void ClearButtons()
        {
            foreach (var b in _buttons)
                b.Visible = false;
        }

        // ── Click routing ─────────────────────────────────────────
        public void HandleClick(int buttonIndex)
        {
            if (_activeView == null) return;
            var descriptors = _activeView.ButtonDescriptors;
            if (buttonIndex < 0 || buttonIndex >= descriptors.Count) return;
            descriptors[buttonIndex].Action?.Invoke();  // null-safe
        }
    }
}"""
)

h2("What ButtonBarController Eliminates from ControlManager")
side_by_side(
    "Current SetVisible — Mixed Concerns  ✗",
"""public void SetVisible(bool isVisible)
{
    visible = isVisible;
    foreach (MedicalControl m in MedicalControls)
    {
        if (isVisible &&
            m == MedicalControls[activeControl])
        {
            m.Visible = true;

            // Button bar sync tangled in
            // visibility management:
            for (int i = 0;
                 i < Buttons.Length; i++)
            {
                if (i >= m.ButtonsText.Length
                    || m.ButtonsText[i] == null)
                    Buttons[i].Visible = false;
                else
                {
                    Buttons[i].Text =
                        m.ButtonsText[i];
                    Buttons[i].Visible = true;
                }
            }
        }
        else
        {
            m.Visible = false;
        }
    }
}""",
    "New SetVisible — Single Concern  ✓",
"""// ModuleHost.SetVisible — ONLY manages
// panel visibility, nothing else:
public void SetVisible(bool isVisible)
{
    _isVisible = isVisible;

    if (!isVisible)
    {
        _navigator.Reset();   // snap to root
        foreach (var v in _views)
            v.Control.Visible = false;
    }
    else
    {
        // Show only the current view
        foreach (var v in _views)
            v.Control.Visible =
                (v == _navigator.Current);
    }
    // Buttons are handled automatically:
    // Navigator.Navigated fires,
    // ButtonBarController.OnNavigated runs,
    // SyncButtons() updates the button bar.
    // Zero button code here.
}"""
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8: COMPONENT 4 — ModuleHost
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
component_header(4, "ModuleHost",
                 "The new ControlManager — a thin coordinator with 5 real responsibilities removed",
                 "1F3964")

h2("What ModuleHost Is")
body(
    "ModuleHost is what ControlManager becomes once its responsibilities have been "
    "redistributed. It is the owner and lifetime manager of a single module's "
    "views and navigator. It has exactly two jobs:"
)
bullet("Own and initialise the views and navigator for one module")
bullet("Manage the visibility of that module's panel at the shell level")
body(
    "Everything else — button sync, click routing, navigation state — is handled by the "
    "components it composes. It wires them together and stays out of their way.",
    italic=True
)

code_block(
"""namespace prgPMR.Core
{
    /// <summary>
    /// Owns the navigator and views for a single medical module.
    /// Replaces ControlManager. Has no knowledge of button bars,
    /// concrete control types, or module-type enums.
    /// </summary>
    public class ModuleHost
    {
        private readonly ModuleNavigator               _navigator;
        private readonly IReadOnlyList<IModuleView>    _views;
        private          bool                          _isVisible;

        // ── Read-only state ───────────────────────────────────────
        public bool             IsVisible  => _isVisible;
        public ModuleNavigator  Navigator  => _navigator;
        public string           ModuleId   { get; }

        // ── Construction ──────────────────────────────────────────
        public ModuleHost(IMedicalModule module)
        {
            ModuleId   = module.ModuleId;
            _navigator = new ModuleNavigator();

            // Module creates its own views, injecting the navigator
            _views = module.CreateViews(_navigator);
        }

        // ── Panel control access (for MainForm to add to pnlMain) ─
        public IEnumerable<UserControl> GetControls() =>
            _views.Select(v => v.Control);

        // ── Visibility ────────────────────────────────────────────
        public void SetVisible(bool isVisible)
        {
            _isVisible = isVisible;

            if (!isVisible)
            {
                _navigator.Reset();            // unwind stack, reload root
                foreach (var v in _views)
                    v.Control.Visible = false;
            }
            else
            {
                // Only the current (top of stack) view is visible
                foreach (var v in _views)
                    v.Control.Visible =
                        (v == _navigator.Current);

                // ButtonBarController sees Navigator.Navigated
                // and syncs the button bar automatically
            }
        }
    }
}"""
)

h2("Line Count Comparison")
make_table(
    ["Metric", "ControlManager (current)", "ModuleHost (redesign)"],
    [
        ["Lines of code",           "~141",                      "~50"],
        ["External dependencies",   "MainForm, all 10 control types, Button[]", "IMedicalModule only"],
        ["Responsibilities",        "7  (factory, nav, visibility, buttons, routing, lifecycle, identity)", "2  (own views, manage visibility)"],
        ["Testable without UI?",    "No — needs WinForms controls and MainForm", "Yes — ModuleNavigator is pure C#"],
        ["Adding a new module",     "Edit this file (switch statement)",         "Never touch this file"],
    ],
    col_widths=[1.8, 2.5, 2.4]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9: COMPONENT 5 — ModuleCoordinator
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
component_header(5, "ModuleCoordinator",
                 "Absorbs module routing from MainForm — makes the shell truly dumb",
                 "1F3964")

h2("The Problem with MainForm Today")
body(
    "MainForm currently manages the active module string, routes button bar clicks "
    "to the correct ControlManager, and loops over all managers to show/hide them. "
    "This is coordination logic that doesn't belong in a WinForms shell. "
    "ModuleCoordinator extracts it, making MainForm a pure host."
)

code_block(
"""namespace prgPMR.Core
{
    /// <summary>
    /// Manages the set of active modules and the shared button bar.
    /// MainForm delegates all routing to this class and becomes a pure
    /// layout host.
    /// </summary>
    public class ModuleCoordinator
    {
        private readonly Dictionary<string, ModuleHost> _hosts = new();
        private readonly ButtonBarController             _buttonBar;
        private          ModuleHost?                     _activeHost;

        public ModuleCoordinator(Button[] sharedButtonBar)
        {
            _buttonBar = new ButtonBarController(sharedButtonBar);
        }

        // ── Module registration ───────────────────────────────────
        public ModuleHost Register(IMedicalModule module)
        {
            var host = new ModuleHost(module);
            _hosts[module.ModuleId] = host;
            return host;                    // caller adds controls to panel
        }

        // ── Module activation ─────────────────────────────────────
        public void Activate(string moduleId)
        {
            if (_activeHost != null)
            {
                _buttonBar.DetachFrom(_activeHost.Navigator);
                _activeHost.SetVisible(false);
            }

            _activeHost = _hosts[moduleId];
            _buttonBar.AttachTo(_activeHost.Navigator);
            _activeHost.SetVisible(true);
        }

        // ── Button click routing ───────────────────────────────────
        public void HandleButtonClick(int buttonIndex) =>
            _buttonBar.HandleClick(buttonIndex);

        // ── Query ─────────────────────────────────────────────────
        public IEnumerable<ModuleHost> AllHosts => _hosts.Values;
    }
}"""
)

h2("How MainForm Simplifies")
side_by_side(
    "MainForm — Before  ✗",
"""public MainForm()
{
    InitializeComponent();
    Button[] bar = [btn0,btn1,
        btn2,btn3,btn4,btn5];
    // ... tag all buttons ...
    ControlManagerDict = [];
    foreach (MedicalControlType type
        in Enum.GetValues<MedicalControlType>())
    {
        ControlManager m = new(type, bar);
        ControlManagerDict.Add(type, m);
        foreach (MedicalControl c
            in m.MedicalControls)
            pnlMain.Controls.Add(c);
    }
    DisplayMedicalControl(Default);
}

private void DisplayMedicalControl(
    MedicalControlType choice)
{
    foreach ((var key, var value)
        in ControlManagerDict)
    {
        value.SetVisible(choice == key);
    }
}

private void ButtonBar_Click(
    object sender, EventArgs e)
{
    if (sender is not Button b
        || b.Tag is not int idx)
        return;
    ControlManagerDict[ActiveMedicalControl]
        .ButtonClicked(idx);
}""",
    "MainForm — After  ✓",
"""public MainForm(
    ModuleRegistry registry,
    Button[] buttonBar)
{
    InitializeComponent();

    _coordinator =
        new ModuleCoordinator(buttonBar);

    foreach (var module in registry.All)
    {
        var host =
            _coordinator.Register(module);
        foreach (var ctrl
            in host.GetControls())
            pnlMain.Controls.Add(ctrl);
    }

    _coordinator.Activate("default");
}

// Nav panel button click:
private void NavButton_Click(
    object sender, EventArgs e)
{
    if (sender is Button b
        && b.Tag is string moduleId)
        _coordinator.Activate(moduleId);
}

// Bottom action button click:
private void ActionButton_Click(
    object sender, EventArgs e)
{
    if (sender is Button b
        && b.Tag is int idx)
        _coordinator.HandleButtonClick(idx);
}"""
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10: END-TO-END WIRING
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("10. End-to-End Wiring — Immunization Module")

body(
    "This section shows how all five components connect together for a real module, "
    "tracing the full lifecycle from startup through adding a new immunization record."
)

h2("IMedicalModule Implementation")
code_block(
"""namespace prgPMR.Modules.Immunization
{
    public class ImmunizationModule : IMedicalModule
    {
        public string ModuleId    => "immunization";
        public string DisplayName => "Immunizations";
        public Version Version    => new(1, 0, 0);

        public IReadOnlyList<IModuleView> CreateViews(IModuleNavigator navigator)
        {
            // Create both views, wire them via the shared navigator
            var listView   = new ImmunizationControl(navigator);
            var detailView = new ImmunizationDetailControl(navigator);

            // Give the list view a reference to the detail view
            // (so it can push it without knowing about ControlManager)
            listView.DetailView = detailView;

            return [listView, detailView];
        }
    }
}"""
)

h2("Startup Sequence")
code_block(
"""// pgmPMRMain.cs
static void Main()
{
    Application.EnableVisualStyles();
    Application.SetCompatibleTextRenderingDefault(false);

    var registry = new ModuleRegistry();
    registry.Register(new ImmunizationModule());   // ← one line per module
    registry.Register(new DoctorVisitModule());
    registry.Register(new MedicationsModule());
    // Adding new modules: just add one line here, nothing else changes

    Application.Run(new LoginForm(registry));
}

// LoginForm.OnLoginSuccess():
var mainForm = new MainForm(registry, sharedButtonBar);
mainForm.Show();"""
)

h2("Sequence: User Adds an Immunization Record")
code_block(
"""USER                    MainForm         ModuleCoordinator   ButtonBarController
 │                           │                    │                    │
 │ click [Immunizations]     │                    │                    │
 │──────────────────────────▶│                    │                    │
 │                           │ NavButton_Click     │                    │
 │                           │────────────────────▶│                    │
 │                           │                     │ Activate("imm...")  │
 │                           │                     │──────────────────▶ │
 │                           │                     │                    │ AttachTo(navigator)
 │                           │                     │                    │ navigator fires Navigated
 │                           │                     │                    │ SyncButtons() runs
 │  [Add | _ | _ | Reset]    │ button bar updated  │                    │
 │◀──────────────────────────│                     │                    │
 │                           │                     │                    │
 │ click [Add]               │                     │                    │
 │──────────────────────────▶│                     │                    │
 │                           │ ActionButton_Click  │                    │
 │                           │────────────────────▶│                    │
 │                           │                     │ HandleButtonClick(0)│
 │                           │                     │────────────────────▶│
 │                           │                     │                    │ descriptors[0].Action()
 │                           │                     │                    │ = ImmunizationControl.Add()
 │                           │                     │                    │
 │                           │  ImmunizationControl.Add():             │
 │                           │  Navigator.Push(detailView,             │
 │                           │      ImmunizationPayload{Mode:Add})     │
 │                           │                     │                    │
 │                           │  ModuleNavigator:                       │
 │                           │  listView.OnDeactivated()               │
 │                           │  detailView.OnActivated(payload)        │
 │                           │    └─ SetButtons(ButtonPresets.DetailAdd)│
 │                           │       └─ ButtonsChanged event fires     │
 │                           │                     │                    │ OnButtonsChanged()
 │                           │                     │                    │ SyncButtons()
 │  [_ | _ | _ | Reset | Save | Cancel]           │                    │
 │◀──────────────────────────│                     │                    │
 │                           │  Navigated event fires                  │
 │                           │  ModuleHost sees it, updates visibility │
 │  [detail form visible]    │                     │                    │
 │                           │                     │                    │
 │ fill form, click [Save]   │                     │                    │
 │──────────────────────────▶│ ActionButton_Click  │                    │
 │                           │────────────────────▶│ HandleButtonClick(4)│
 │                           │                     │────────────────────▶│
 │                           │                     │                    │ descriptors[4].Action()
 │                           │                     │                    │ = detailView.Save()
 │                           │  detailView.Save():                     │
 │                           │  await _repo.AddAsync(record)           │
 │                           │  Navigator.Pop()                        │
 │                           │                     │                    │
 │                           │  ModuleNavigator:                       │
 │                           │  detailView.OnDeactivated()             │
 │                           │  listView.OnReactivated()               │
 │                           │    └─ InitializeGrid() reloads from DB  │
 │                           │       └─ ButtonsChanged fires           │
 │                           │                     │                    │ SyncButtons()
 │  [Add | _ | _ | Reset]   │ button bar restored  │                    │
 │  [grid refreshed]         │                     │                    │"""
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11: FULL COMPARISON
# ══════════════════════════════════════════════════════════════════════════════
doc.add_page_break()
h1("11. Old Design vs New Design — Full Comparison")

make_table(
    ["Concern", "Current Design", "Redesign"],
    [
        ["Module types",          "Hardcoded enum in MainForm, switch in ControlManager",           "Self-registering IMedicalModule, no enum"],
        ["Navigation model",      "int activeControl — index into a List",                          "Stack<IModuleView> — push/pop"],
        ["Navigate forward",      "Manager.NextControl(DataInterface?) — null = reset",             "Navigator.Push(view, payload) — explicit"],
        ["Navigate back",         "Manager.PreviousControl() — index decrement",                    "Navigator.Pop() — always safe, fires lifecycle"],
        ["Button bar data",       "string[] + Action[] parallel arrays (can misalign)",              "IReadOnlyList<ButtonDescriptor> — one object"],
        ["Button bar sync",       "SetVisible() loop — navigation + buttons tangled",               "ButtonBarController.SyncButtons() — dedicated"],
        ["Button click routing",  "ControlManager.ButtonClicked(int)",                              "ButtonBarController.HandleClick(int) — null-safe"],
        ["Button presets",        "Two dictionaries duplicated in every concrete control",           "ButtonPresets static class — one place"],
        ["Back-reference",        "MedicalControl.Manager calls RefreshVisibility()",               "MedicalControl raises ButtonsChanged event"],
        ["Circular dependency",   "ControlManager ↔ MedicalControl (mutual)",                       "Broken: views raise events, coordinator listens"],
        ["Adding a module",       "Edit MainForm enum + ControlManager switch + create controls",    "New class, one registration line"],
        ["Reset on hide",         "Not implemented — returns to whatever state was left",            "Navigator.Reset() in SetVisible(false)"],
        ["Testability",           "Cannot test ControlManager without all WinForms controls",        "ModuleNavigator is pure C# — no WinForms needed"],
        ["MainForm knowledge",    "Owns MedicalControlType enum + active module tracking",          "Knows nothing — delegates to ModuleCoordinator"],
        ["Button bar ownership",  "public Button[] Buttons on ControlManager",                      "private Button[] inside ButtonBarController"],
    ],
    col_widths=[1.65, 2.3, 2.3],
    row_colors=[
        "FCEAEA","FCEAEA","FCEAEA","FCEAEA",
        "FCEAEA","FCEAEA","FCEAEA","FCEAEA",
        "FCEAEA","FCEAEA","FCEAEA","FCEAEA",
        "FCEAEA","FCEAEA","FCEAEA",
    ]
)
body("Red rows indicate current problems. Each maps directly to a component in the redesign.", italic=True, color=GRAY)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 12: MIGRATION PATH
# ══════════════════════════════════════════════════════════════════════════════
h1("12. Migration Path — Getting There Without Breaking Everything")

body(
    "A redesign of this scale does not have to be done all at once. "
    "The following sequence migrates the system component by component, keeping "
    "the application runnable at every step."
)

make_table(
    ["Step", "Action", "Keeps Working?", "Risk"],
    [
        ["1", "Add ButtonDescriptor record and ButtonPresets static class — no existing code changes",         "Yes", "None"],
        ["2", "Add IModuleView interface and IModuleNavigator interface — not yet used",                       "Yes", "None"],
        ["3", "Add ModuleNavigator class — not yet wired in",                                                 "Yes", "None"],
        ["4", "Add ButtonBarController class — not yet wired in",                                             "Yes", "None"],
        ["5", "Migrate ImmunizationControl to implement IModuleView, using ButtonPresets",                    "Yes", "Low — one module"],
        ["6", "Migrate ImmunizationDetailControl — same approach",                                           "Yes", "Low"],
        ["7", "Create ImmunizationModule — introduces CreateViews()",                                        "Yes", "Low"],
        ["8", "Create ModuleHost — use it for Immunization only, keep old ControlManager for others",         "Yes", "Medium"],
        ["9", "Create ModuleCoordinator — wire Immunization through it, others still use old path",          "Yes", "Medium"],
        ["10","Migrate remaining modules one by one using the same pattern",                                   "Yes", "Low per step"],
        ["11","Remove ControlManager, MedicalControlType enum, and old switch statement",                     "Yes", "Low — final cleanup"],
    ],
    col_widths=[0.4, 3.5, 1.1, 0.9]
)

insight(
    "Steps 1–4 are additive — no existing code is touched. The old system stays fully "
    "functional while the new infrastructure is built alongside it. Only from step 5 "
    "onward does any migration actually occur, and it can be done one module at a time."
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 13: SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h1("13. Summary — What the Redesign Achieves")

make_table(
    ["Principle", "How the Redesign Addresses It"],
    [
        ["Single Responsibility",  "Each component has exactly one job. ModuleHost owns views. Navigator manages state. ButtonBarController manages buttons."],
        ["Open/Closed",            "New modules require zero changes to any existing file. Register one class, done."],
        ["Liskov Substitution",    "IModuleNavigator allows fake navigators in tests. IModuleView allows any UserControl to be a view."],
        ["Interface Segregation",  "Views depend on IModuleNavigator (not ModuleNavigator). Tests provide a fake. No WinForms dependency needed."],
        ["Dependency Inversion",   "Views depend on IModuleNavigator (abstraction). ButtonBarController depends on IModuleView (abstraction). Nothing depends on concrete classes."],
        ["Observer Pattern",       "ButtonsChanged event decouples view state from button rendering. No polling, no callbacks into the owner."],
        ["Stack Navigation",       "Any module can have as many view levels as needed without any change to the navigation infrastructure."],
        ["Testability",            "ModuleNavigator has zero WinForms dependencies — test push/pop/reset with plain xUnit, no mocking framework required."],
    ],
    col_widths=[1.7, 4.8]
)

body(
    "The core idea ControlManager had — a coordinator that knows about a module's views and "
    "manages their lifecycle — was always right. The redesign preserves that idea and builds "
    "it properly: the navigation state machine is explicit and testable, the button bar is "
    "event-driven and decoupled, the views own their lifecycle, and the shell knows nothing "
    "about any of it.",
    italic=True
)

# ── Save ─────────────────────────────────────────────────────────────────────
output_path = "E:/Dev/GitHub/medical-dashboard/ControlManager_Redesign.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
